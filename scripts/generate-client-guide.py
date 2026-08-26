#!/usr/bin/env python3
"""Generate the illustrated Blackout Ridge client, facilitator, and new-player guide."""

from datetime import date
from io import BytesIO
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "client-guide"
SHOTS = OUT_DIR / "screenshots"
OUT_FILE = OUT_DIR / "Blackout_Ridge_Complete_Client_Guide.docx"
LOGO = ROOT / "prototype" / "assets" / "images" / "logo2.webp"
TODAY = date.today().strftime("%d %B %Y")
GAME_URL = "http://151.243.217.9:4173/"

RED = "D7352C"
DARK_RED = "8E211C"
INK = "172126"
SLATE = "4D5F66"
PALE = "F2F5F5"
PALE_RED = "FCEDEC"
PALE_GREEN = "EAF6EF"
PALE_AMBER = "FFF5DE"
WHITE = "FFFFFF"
GREEN = "25734B"
AMBER = "9C6516"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_header(section) -> None:
    section.header.is_linked_to_previous = False
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("BLACKOUT RIDGE  /  COMPLETE CLIENT GUIDE")
    run.bold = True
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(RED)

    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BLACKOUT RIDGE  •  ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def docx_image(path: Path):
    """Return an OOXML-compatible image stream for WebP source assets."""
    if path.suffix.lower() != ".webp":
        return str(path)
    stream = BytesIO()
    with Image.open(path) as source:
        source.save(stream, format="PNG", optimize=True)
    stream.seek(0)
    return stream


def add_hyperlink(paragraph, text: str, url: str, color=RED):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def configure_document(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Blackout Ridge — Complete Client, Facilitator & New-Player Guide"
    props.subject = "Operating, facilitating, playing, debriefing, and troubleshooting Blackout Ridge"
    props.author = "Blackout Ridge delivery team"
    props.keywords = "Blackout Ridge, client guide, facilitator, participant, multiplayer, workshop"
    props.comments = "Generated from the implemented runtime and current game screens."

    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(16)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(17)
    section.right_margin = Mm(17)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(8)
    set_repeat_header(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(INK)

    for name, size, color in (
        ("Heading 1", 20, RED),
        ("Heading 2", 14, DARK_RED),
        ("Heading 3", 11.5, INK),
    ):
        style = styles[name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(11)
        style.paragraph_format.space_after = Pt(5)

    styles["List Bullet"].font.name = "Aptos"
    styles["List Bullet"].font.size = Pt(10)
    styles["List Number"].font.name = "Aptos"
    styles["List Number"].font.size = Pt(10)


def page_break(doc: Document) -> None:
    doc.add_page_break()


def h1(doc: Document, text: str, subtitle: str | None = None) -> None:
    doc.add_heading(text, level=1)
    if subtitle:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r = p.add_run(subtitle)
        r.italic = True
        r.font.color.rgb = RGBColor.from_string(SLATE)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead):])
    else:
        p.add_run(text)


def bullets(doc: Document, items, numbered=False) -> None:
    style = "Normal" if numbered else "List Bullet"
    for index, item in enumerate(items, 1):
        p = doc.add_paragraph(style=style)
        if numbered:
            p.paragraph_format.left_indent = Mm(6)
            p.paragraph_format.first_line_indent = Mm(-6)
        p.paragraph_format.space_after = Pt(3)
        if numbered:
            r = p.add_run(f"{index}.  ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(RED)
        if isinstance(item, tuple):
            lead, body = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(body)
        else:
            p.add_run(item)


def table(doc: Document, headers, rows, widths=None, small=False) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    t.autofit = widths is None
    hdr = t.rows[0]
    set_repeat_table_header(hdr)
    for i, value in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(value))
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.4 if small else 9)
        if widths:
            cell.width = widths[i]
    for row_index, values in enumerate(rows):
        row = t.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            if row_index % 2:
                set_cell_shading(cell, PALE)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            r.font.size = Pt(8.3 if small else 9)
            if widths:
                cell.width = widths[i]
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def callout(doc: Document, title: str, body: str, tone="info") -> None:
    fill, accent = {
        "info": (PALE, SLATE),
        "warning": (PALE_AMBER, AMBER),
        "danger": (PALE_RED, DARK_RED),
        "success": (PALE_GREEN, GREEN),
    }[tone]
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Mm(4)
    t.columns[1].width = Mm(166)
    accent_cell, content = t.rows[0].cells
    set_cell_shading(accent_cell, accent)
    set_cell_shading(content, fill)
    set_cell_margins(accent_cell, 0, 0, 0, 0)
    set_cell_margins(content, 110, 150, 110, 150)
    p = content.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(accent)
    p = content.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    p.runs[0].font.size = Pt(9.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


FIGURE = 0


def screenshot(doc: Document, filename: str, caption: str, notes: list[str] | None = None, width=Inches(6.55)) -> None:
    global FIGURE
    path = SHOTS / filename
    if not path.exists():
        callout(doc, "Screenshot missing", f"Expected documentation image: {filename}", "danger")
        return
    FIGURE += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    run.add_picture(docx_image(path), width=width)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.keep_with_next = bool(notes)
    cap.paragraph_format.space_after = Pt(4)
    r = cap.add_run(f"Figure {FIGURE}. {caption}")
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    if notes:
        t = doc.add_table(rows=1, cols=len(notes))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, note in enumerate(notes):
            cell = t.rows[0].cells[i]
            set_cell_shading(cell, PALE)
            set_cell_margins(cell, 70, 90, 70, 90)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(f"{i + 1}  ")
            r.bold = True
            r.font.color.rgb = RGBColor.from_string(RED)
            r = p.add_run(note)
            r.font.size = Pt(8)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)


def two_screens(doc: Document, left_name: str, left_caption: str, right_name: str, right_caption: str) -> None:
    global FIGURE
    t = doc.add_table(rows=2, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for col, (name, caption) in enumerate(((left_name, left_caption), (right_name, right_caption))):
        path = SHOTS / name
        p = t.rows[0].cells[col].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if path.exists():
            p.add_run().add_picture(docx_image(path), width=Inches(2.35))
        FIGURE += 1
        p = t.rows[1].cells[col].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"Figure {FIGURE}. {caption}")
        r.italic = True
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_paragraph()


def chapter(doc: Document, number: str, title: str, subtitle: str) -> None:
    page_break(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"SECTION {number}")
    r.bold = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(RED)
    h1(doc, title, subtitle)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(docx_image(LOGO), width=Inches(3.0))
    doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Complete Client, Facilitator\n& New-Player Guide")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("How to prepare, host, play, explain, debrief, and support every Blackout Ridge session")
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(SLATE)
    doc.add_paragraph()
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    set_cell_shading(cell, INK)
    set_cell_margins(cell, 160, 220, 160, 220)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CLIENT DELIVERY EDITION  •  VERSION 1.0")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(WHITE)
    r.font.size = Pt(11)
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"Prepared {TODAY}")
    r.font.color.rgb = RGBColor.from_string("C9D4D8")
    r.font.size = Pt(9)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Current hosted build: ").bold = True
    add_hyperlink(p, GAME_URL, GAME_URL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Facilitator material includes story spoilers. Give participants only Sections 2–4 unless spoilers are acceptable.")
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string(DARK_RED)
    screenshot(doc, "01-home-boot.webp", "The in-world station boot screen participants see on entry.", width=Inches(5.8))


def build() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    cover(doc)

    page_break(doc)
    h1(doc, "Document control")
    table(doc, ["Item", "Value"], [
        ["Document", "Blackout Ridge — Complete Client, Facilitator & New-Player Guide"],
        ["Version", "1.0 / client delivery edition"],
        ["Prepared", TODAY],
        ["Runtime", "Browser-based synchronized multiplayer game; Node server on port 4173"],
        ["Supported crew sizes", "2, 4, 5, 6, or 7 participants, plus one facilitator"],
        ["Recommended duration", "Approximately 70 minutes: 5-minute briefing, 43-minute play, 20–25-minute debrief"],
        ["Primary audience", "Client sponsors, workshop producers, facilitators, technical hosts, and new participants"],
    ])
    callout(doc, "Security and privacy", "Facilitator and player links contain bearer tokens. Treat them like private access links. The room-facing shared link is public/read-only. Crew Chat and debrief notes are stored with the session and appear in the facilitator export.", "warning")
    h2(doc, "How to use this guide")
    table(doc, ["Reader", "Read first", "Then use"], [
        ["Client sponsor", "Sections 1 and 8", "Sections 5–7 for the full learning design"],
        ["Technical host", "Sections 2 and 9", "Appendix A preflight checklist"],
        ["Facilitator", "Sections 2, 5, 6, and 7", "Appendix B scripts and Appendix C recovery table"],
        ["New participant", "Sections 3 and 4", "Avoid Section 6 if you want a spoiler-free experience"],
        ["Observer / evaluator", "Sections 1, 5, 7, and 8", "Appendix D playtest evidence prompts"],
    ])
    h2(doc, "Contents")
    table(doc, ["Section", "Purpose"], [
        ["1. What Blackout Ridge is", "Product intent, audience, learning model, and experience architecture"],
        ["2. Quick start and technical setup", "Launch, host, session creation, links, devices, and preflight"],
        ["3. New-player guide", "Join, claim a role, read evidence, communicate, recommend, and act"],
        ["4. Screen-by-screen reference", "Shared display, private terminal, facilitator console, chat, telemetry"],
        ["5. Facilitator runbook", "Care briefing, opening script, timing, controls, and room management"],
        ["6. Seven-stage walkthrough (spoilers)", "Exact objective, decisions, facilitation notes, and consequences"],
        ["7. Endings and debrief", "Ending logic, six-step debrief, first steps, and export"],
        ["8. Explaining the game to a client", "Value proposition, demo narrative, evidence, and success criteria"],
        ["9. Troubleshooting and recovery", "Network, browser, token, role, timer, mobile, and audio issues"],
        ["Appendices", "Checklists, scripts, controls, glossary, and one-page quick reference"],
    ])

    chapter(doc, "1", "What Blackout Ridge is", "A synchronized leadership simulation about challenging calm automated labels when human reality is incomplete.")
    h2(doc, "1.1 Experience in one paragraph")
    para(doc, "Blackout Ridge is a live-facilitated, browser-based survival-horror leadership simulation. A storm has damaged a remote relay station, a colleague named Mara Venn is missing, and the station is producing plausible operational conclusions from incomplete data. Every participant receives a different private evidence lens. The group must speak, combine evidence, make seven timed decisions, manage both physical danger and the official record, and then connect the experience to real AI-enabled workplace decisions.")
    h2(doc, "1.2 What the game teaches")
    bullets(doc, [
        ("Challenge the label, not only the number. ", "A clean status can measure the wrong object. OPERATIONAL describes relay transmission, not human safety."),
        ("Ask what evidence is missing. ", "Players see different fragments and must reconstruct provenance, human context, dependencies, and guardrails."),
        ("Act under uncertainty without pretending certainty. ", "Every stage demands one bounded action; endless analysis allows defaults to harden."),
        ("Keep people visible in system records. ", "Mara and the former Lark Shift expose the danger of replacing named people with occupancy or status abstractions."),
        ("Make corrections actionable to receivers. ", "A warning succeeds only when an outside authority understands the invalid status, person, location, hazard, evidence, and requested action."),
    ])
    h2(doc, "1.3 What it is not")
    bullets(doc, [
        "It is not a technical exam; roles are leadership perspectives, not professional certifications.",
        "It is not a hidden-information competition; players may read or paraphrase their private evidence.",
        "It is not a psychometric assessment; people receive no numerical score or personality diagnosis.",
        "It is not seven disconnected puzzles; decisions persist and change later evidence, safety, trust, and endings.",
        "It is not a generative-AI adjudicator; outcome logic is deterministic and testable.",
    ])
    h2(doc, "1.4 Four coordinated views")
    table(doc, ["View", "Audience", "Purpose", "Contains private information?"], [
        ["Game start / join", "Host and participants", "Create a shift, enter a code, claim an available responsibility", "No clues yet"],
        ["Shared station", "Room display or screen share", "Common incident, clock, telemetry, evidence ledger, roster, consequence", "No"],
        ["Private role terminal", "One participant", "Role-specific evidence, recommendation, intervention, lead decision tools", "Yes"],
        ["Facilitator console", "Facilitator only", "Pacing, decisions, all reports, safety, recovery, ending preview, debrief", "Yes—keep private"],
    ], small=True)
    callout(doc, "Core loop", "Story pressure → Private evidence → Crew discussion → Decision preview → Commitment → Cause-and-effect reveal. Every stage begins from the previous consequence, so the crew always knows why the new emergency is happening and what its current decision may preserve or put at risk.", "success")
    h2(doc, "1.5 The escalating story arc")
    table(doc, ["Act", "Stages", "Escalation"], [
        ["Doubt the handover", "1–2", "Mara is absent; weak human evidence conflicts with calm defaults; lightning then forces the crew to sacrifice future capabilities."],
        ["Follow the contradiction", "3–5", "The disputed signal points away from the road, an exposed field test locates a route beneath the station, and the crew discovers Mara inside a legacy hold."],
        ["Correct and survive", "6–7", "The crew must replace the false outside record with actionable truth while simultaneously controlling the interlock, carrier, Mara's rescue, and its own air."],
    ], small=True)

    chapter(doc, "2", "Quick start and technical setup", "Everything a host needs to move from a stopped server to a ready room.")
    h2(doc, "2.1 Minimum setup")
    table(doc, ["Item", "Minimum", "Recommended"], [
        ["Facilitator", "1 laptop/browser", "Laptop plus a second monitor"],
        ["Participants", "One browser per person", "Laptop or tablet; phone supported when necessary"],
        ["Shared display", "Optional", "TV/projector or a dedicated screen-shared tab"],
        ["Voice", "Same room or conference call", "Teams/Zoom plus headphones where useful"],
        ["Browser", "Current Chrome/Edge/Firefox/Safari", "Chrome or Edge on a stable connection"],
        ["Server", "Node.js runtime, port 4173", "Wired host, persistent terminal, firewall rule for the chosen network"],
    ])
    h2(doc, "2.2 Start the application")
    bullets(doc, [
        "Open a terminal in the project folder.",
        "Run: npm run dev",
        "Wait for: Blackout Ridge running at http://localhost:4173",
        "On the host machine, open http://localhost:4173/",
        f"For the current hosted build, open {GAME_URL}",
    ], numbered=True)
    callout(doc, "Network note", "localhost works only on the server machine. Other devices need a reachable LAN or public address and access to TCP port 4173. If a link opens on the host but not on participant devices, check routing/firewall before changing the game.", "warning")
    h2(doc, "2.3 Create a shift")
    screenshot(doc, "02-host-config.webp", "Command Authority chooses the number of participant terminals before arming the session.", ["Choose 2, 4, 5, 6, or 7 players.", "The facilitator becomes Game Master; they are not counted as a player.", "Arm the shift to create the session code and private facilitator authority."])
    bullets(doc, [
        "Engage the station from the opening screen.",
        "Choose Command Authority / host mode.",
        "Select the exact number of participants who will play roles.",
        "Arm the shift. The application creates a six-character session code and private facilitator token.",
        "Keep the facilitator tab private. Copy the crew link and open the Room Display in a separate tab.",
    ], numbered=True)
    h2(doc, "2.4 Assemble the room")
    screenshot(doc, "04-assembly-lobby.webp", "The facilitator assembly lobby shows connected terminals, the participant-care gate, crew link, room display, and start control.", ["Roster: each required terminal must be connected.", "Participant care: the facilitator must read and confirm the briefing.", "Start remains locked until both conditions are true."])
    h3(doc, "Link handling")
    table(doc, ["Link", "Who receives it", "Rule"], [
        ["Crew link", "All participants", "Safe to share with the intended cohort; players claim one open role"],
        ["Player link", "Only that participant", "Contains a bearer token; do not paste into public chat or screenshots"],
        ["Shared display", "Room display / screen share", "Read-only and intentionally excludes private evidence and Crew Chat"],
        ["Facilitator link", "Facilitator only", "Contains full control and all recommendations; treat as confidential"],
    ])
    h2(doc, "2.5 Five-minute preflight")
    bullets(doc, [
        "Open the shared display at 100% browser zoom and full-screen it.",
        "Confirm every participant can see their name, role, and private evidence card.",
        "Ask players not to share their browser tab or player URL.",
        "Test voice communication; use Crew Chat as fallback, not as the only plan unless required.",
        "Confirm the facilitator can Pause, add 30 seconds, send a prompt, and activate Safety Pause.",
        "Explain that audio is opt-in and can remain off.",
        "Read the participant-care statement and confirm it in the lobby.",
        "Close unrelated notifications and avoid browser translation or aggressive content blockers.",
    ])

    chapter(doc, "3", "New-player guide", "A spoiler-light explanation participants can follow without learning the answers.")
    h2(doc, "3.1 Join and claim a responsibility")
    screenshot(doc, "03-join-flow.webp", "A new participant enters a callsign and claims an available role terminal.", ["Enter the name others should recognize.", "Read the participant-care notice before connecting.", "Only unclaimed roles can be selected; one browser becomes one private terminal."])
    bullets(doc, [
        "Open the crew link from the facilitator.",
        "Enter your display name or workshop callsign.",
        "Choose one available responsibility. If a role disappears, another participant claimed it first.",
        "Press Connect Role Terminal and keep that tab open.",
        "If you refresh, return using the same player link or browser; the token restores your role.",
    ], numbered=True)
    h2(doc, "3.2 The responsibilities")
    table(doc, ["Role", "What this role protects", "Once-per-session override"], [
        ["Station Lead", "Accountability, pace, unresolved disagreement, crew commitment", "Hold the Clock"],
        ["Signal Analyst", "Source, timestamps, waveform meaning, preservation and replay", "Recover Fragment"],
        ["Systems Engineer", "Dependencies and what a status actually measures", "Trace Dependency"],
        ["Operations Officer", "Bounded action, route, equipment, timing, abort conditions", "Controlled Test"],
        ["Field Liaison", "Mara, human context, staff practice, people missing from clean data", "Human Check"],
        ["Protocol Officer", "Guardrails, uncertainty, receiver requirements, safety thresholds", "Safety Condition"],
        ["Comms Officer (7 players)", "Receiver trust, message clarity, acknowledgement and readback", "Request Readback"],
    ], small=True)
    callout(doc, "Combined roles", "Two-player games use Command & Field plus Signal & Systems. Four- and five-player games combine selected responsibilities. The story and all seven decisions remain intact; each combined card includes the necessary evidence lenses.", "info")
    h2(doc, "3.3 What to do in every stage")
    screenshot(doc, "04b-player-lead-stage1.webp", "A player terminal now turns each private brief into four plain-language actions: understand the evidence, interpret its confidence, say what matters aloud, and make a recommendation.", ["WHAT YOU KNOW separates evidence from interpretation.", "WHY IT MATTERS connects that evidence to the current crew decision.", "WHAT TO SAY OUT LOUD gives a direct discussion prompt.", "The consequence preview names what the proposed decision preserves, costs, and still leaves unresolved."])
    bullets(doc, [
        ("Read. ", "Identify what your evidence actually proves and what remains an assumption."),
        ("Speak. ", "Share the evidence aloud or in Crew Chat. State its confidence: confirmed, likely, stale, unknown, or contradicted."),
        ("Recommend. ", "Complete the sentence: Because [evidence], we should [action], otherwise [risk]. Then choose Share with Station Lead. A visible confirmation shows it was routed."),
        ("Question. ", "Ask which source produced a clean label and what that source is authorized to claim."),
        ("Commit. ", "The Station Lead summarizes and submits the crew decision. The facilitator can also resolve when necessary."),
        ("Observe consequence. ", "Read what changed; earlier choices affect later clarity, safety, trust, and options."),
    ])
    h2(doc, "3.4 Recommendations are not public chat")
    table(doc, ["Channel", "Who sees the text", "Best use"], [
        ["Your Recommendation", "You, Station Lead, and facilitator; peers see only that you submitted", "Short stage assessment tied to your private evidence"],
        ["Crew Chat", "All authenticated players and facilitator", "Discussion, remote fallback, coordination, accessibility"],
        ["Voice", "Everyone on the room/call", "Primary discussion channel and fastest synthesis"],
        ["Shared display", "Anyone with shared URL", "Common incident only; cannot read or post chat"],
    ])
    h2(doc, "3.5 Crew Chat")
    screenshot(doc, "09a-crew-chat.webp", "Crew Chat uses stable role colors so participants can distinguish senders quickly.", ["Other roles retain distinct colors by responsibility.", "Your own messages appear red and align right.", "Unread messages produce a badge and temporary alert; the shared display has no chat access."])
    bullets(doc, [
        "Open the lower-right Crew Chat button.",
        "Messages synchronize immediately to authenticated players and the facilitator.",
        "Lead is amber, Signal cyan, Systems violet, Operations green, Field rose, Protocol blue, Comms orange, Facilitator yellow; your own bubble is red.",
        "Messages are limited to 500 characters. The session retains the latest 100 messages.",
        "Do not use chat for sensitive personal disclosures; it is included in the facilitator’s exported record.",
    ])
    h2(doc, "3.6 Mobile use")
    two_screens(doc, "04d-player-lead-mobile.webp", "Structured participant note and decision preview on a phone.", "09b-player-signal-mobile.webp", "Live private Signal terminal on a phone.")
    bullets(doc, [
        "Use portrait orientation for reading and chat; landscape can help with dense stage controls.",
        "Scroll within the page when the current control is below the evidence card.",
        "The chat button stays in the lower-right corner; close chat before using controls underneath it.",
        "Browser zoom should remain at 100%. Increase operating-system text size if accessibility requires it.",
    ])

    chapter(doc, "4", "Screen-by-screen reference", "What each interface is showing and what the user should do with it.")
    h2(doc, "4.1 Shared station")
    screenshot(doc, "04a-shared-stage1.webp", "The shared Stage 1 station view keeps the incident, timer, challenge stack, signal, tracks, evidence, and roster in one common picture.", ["Mission and timer establish the current decision window.", "Incident pressures and three persistent tracks show compounding consequences.", "Verified evidence and roster are public; private cards remain hidden."])
    h3(doc, "Persistent tracks")
    table(doc, ["Track", "Labels", "Interpretation"], [
        ["Upper Air", "Stable → Trace → Unsafe → Critical", "Surface operating safety and remaining pressure margin"],
        ["External Trust", "Unverified → Heard → Credible", "Whether Cordon Control will act on Ridge traffic"],
        ["Official Status", "Drafting → Repeating → Hardening → Filed", "How close a false calm conclusion is to becoming authoritative"],
    ])
    h2(doc, "4.2 Private role terminal")
    screenshot(doc, "08-player-signal-stage3.webp", "Signal Analyst sees a role-specific brief during the degraded-signal stage.", ["Private evidence is deliberately incomplete.", "The confidence label describes evidence quality, not whether the player is correct.", "Archive Role Brief downloads a text copy for remote or accessibility use."])
    h3(doc, "How to read a Participant Note")
    table(doc, ["Block", "Meaning", "Participant action"], [
        ["What you know", "The role's private observation or record", "Separate the evidence from your preferred action"],
        ["Confidence", "Why the evidence is confirmed, likely, stale, contradicted, or unknown", "Say the confidence level when sharing it"],
        ["Why it matters", "The connection to this stage's exact decision", "Name what changes if the crew ignores it"],
        ["What to say out loud", "A direct question or sentence starter", "Use it to bring the evidence into discussion"],
        ["Your recommendation", "A private summary routed to Lead and facilitator", "Link evidence, action, and risk in one concise statement"],
    ], small=True)
    h2(doc, "4.3 Facilitator console")
    screenshot(doc, "04c-facilitator-stage1.webp", "The facilitator console combines live control, all role submissions, decision controls, station state, signal, evidence, and event history.", ["Left: pace, safety, prompts, reports, and recovery.", "Center: current crew decision; facilitator can resolve if needed.", "Right: authoritative state, evidence, signal integrity, and audit history."])
    h3(doc, "Controls and correct use")
    table(doc, ["Control", "Use when", "Do not use to"], [
        ["Pause / Resume", "Technical interruption, room management, necessary clarification", "Remove normal decision pressure"],
        ["+30 sec", "Compensate for a connection or interface delay", "Reward indecision or steer toward a better choice"],
        ["Safety Pause", "A participant needs intensity reduced or time out", "Create dramatic effect"],
        ["Send Prompt", "The group is stuck on the rule or missing question", "Reveal a private answer"],
        ["Mark Absent", "A player steps out; observer mode blocks game actions but leaves chat available", "Punish non-participation"],
        ["Re-send Card", "A player lost or cannot locate their private evidence", "Broadcast another role’s evidence"],
        ["Reassign Role", "A role must move to a connected participant", "Change outcomes after seeing the preview"],
        ["Resolve Decision", "Facilitator is authorized to commit or recover the stage", "Silently replace the crew’s actual choice"],
    ], small=True)
    h2(doc, "4.4 Consequence screens")
    screenshot(doc, "04e-stage1-consequence-mobile.webp", "Every commitment opens a cause-and-effect reveal before the next stage can begin.", ["THE CREW COMMITTED repeats the exact choice.", "WHAT CHANGED NOW names immediate gains and costs.", "Persistent tracks show the updated pressure.", "WHY THE NEXT STAGE HAPPENS provides the narrative bridge into the escalating emergency."] , width=Inches(3.6))
    screenshot(doc, "05-shared-stage2.webp", "Stage 2 shared view: power allocation changes which later systems and evidence paths remain available.", ["Before commitment, the preview separates preserved capabilities from costs and warnings.", "After commitment, the consequence freezes the timer and names what changed.", "Only the facilitator advances after the room can explain the cause-and-effect link."])

    chapter(doc, "5", "Facilitator runbook", "A practical, safe, and repeatable 70-minute delivery sequence.")
    h2(doc, "5.1 Recommended timing")
    table(doc, ["Block", "Time", "Facilitator goal"], [
        ["Arrival and technical check", "Before start", "Links, devices, voice, shared display, private tabs"],
        ["Participant care and opening", "5 minutes", "Consent, off-ramp, premise, interaction loop"],
        ["Stages 1–7", "43 minutes designed clocks", "Keep decisions bounded; let consequences teach"],
        ["Transition", "2 minutes", "Let the ending land; separate system record from human reality"],
        ["Debrief", "20–25 minutes", "Reconstruct, map roles, diagnose hesitation, bridge to work, record first steps"],
    ])
    h2(doc, "5.2 Participant-care script")
    callout(doc, "Read before starting", "“This is a fictional survival-horror leadership exercise involving emergency pressure, missing-person themes and implied danger. If at any point you need to step out of the scenario, message me privately and I will accommodate that without drawing attention to it. You do not need to explain why.”", "warning")
    para(doc, "Only choose Confirm Brief Delivered after reading the statement. The game cannot start until every configured terminal is connected and the briefing is confirmed. Safety Pause freezes the clock and replaces active pressure with a neutral pause message on every screen.")
    h2(doc, "5.3 Opening script")
    callout(doc, "Suggested wording", "“You are the night crew at Blackout Ridge Signal Station. The storm has damaged the mast. Mara Venn is missing. The station is producing calm operational answers, but each of you can see different evidence. Share what matters, name what you are assuming, and make one crew decision before each window closes. You may be unable to save everything. Operational does not mean safe.”", "info")
    h2(doc, "5.4 Stage rhythm")
    bullets(doc, [
        "Open with one sentence naming the current incident; stop talking within 30 seconds.",
        "Give participants time to read their cards silently.",
        "Listen for missing roles and ask, “Whose evidence have we not heard?” rather than asking for a specific clue.",
        "At midpoint, name remaining time. At 30 seconds, ask the Station Lead to summarize and commit.",
        "After commitment, read the consequence and ask one short noticing question. Do not debrief the full lesson mid-game.",
        "Advance only when the group understands what changed operationally.",
    ], numbered=True)
    h2(doc, "5.5 Facilitation boundaries")
    table(doc, ["Do", "Avoid"], [
        ["Ask what the output measures", "Explain the central OPERATIONAL twist during play"],
        ["Ask which role has not spoken", "Read private cards for connected players"],
        ["Protect participant choice and safety", "Secretly alter the ending to make it nicer"],
        ["Correct a proven interface/transcription error with an audit note", "Override a choice because the group chose poorly"],
        ["Let weak choices fail forward", "Restart a stage to chase an ideal route"],
        ["Keep the shared and human stakes separate", "Reduce the result to winning, losing, or points"],
    ])
    h2(doc, "5.6 Before advancing to the next stage")
    bullets(doc, [
        "The committed decision shown on screen matches what the Station Lead intended.",
        "The room has read the consequence headline and can name the operational change.",
        "Any timer or connection loss has been handled transparently.",
        "No participant-care request is waiting for a private response.",
        "You have not explained the next reveal or interpreted the group’s behavior for them.",
        "The shared display, player terminals, and facilitator console all show the resolved state.",
    ])

    chapter(doc, "6", "Seven-stage walkthrough", "FACILITATOR-ONLY SPOILERS: exact decisions, intended joins, risks, and stage-to-stage handoffs.")
    callout(doc, "Spoiler boundary", "Do not give this section to first-time participants unless the client explicitly wants an open-book demonstration. The learning depends on discovering that several calm labels describe system proxies rather than human reality.", "danger")
    h2(doc, "Stage overview")
    table(doc, ["Stage", "Time", "Question", "Required crew decision"], [
        ["1. Challenge the handover", "4 min", "Which defaults need human review?", "Hold exactly 3 of 5 proposed statements"],
        ["2. Keep the station alive", "5 min", "Which future capabilities survive?", "Power exactly 3 of 6 optional circuits"],
        ["3. Decide what to trust", "6 min", "Which traffic deserves scarce reconstruction?", "Allocate exactly 3 signal slots"],
        ["4. Test Lower Route", "6 min", "Is Lower Route the flooded road?", "Runner, destination, abort rule; then 3 moves"],
        ["5. Open the buried station", "7 min", "How can Mara be verified without trapping more people?", "Hatch procedure, respirator use, watcher"],
        ["6. Build the truth", "6 min", "Can Cordon act on a verified correction?", "Five-part correction; possible shorter retry"],
        ["7. Survive the correction", "9 min", "Can the crew control physical and official truth together?", "Two rounds across Lock, Signal, People"],
    ], small=True)

    h2(doc, "6.1 Stage 1 — Challenge the handover")
    screenshot(doc, "04c-facilitator-stage1.webp", "Stage 1 facilitator view with five plausible defaults and exactly three available audit holds.")
    bullets(doc, [
        ("Situation. ", "The outgoing crew left hurriedly. Mara is absent. The station calmly fills gaps with five defaults."),
        ("Decision. ", "Hold exactly three: Local Service Fault; No Wider Risk Confirmed; Venn on Lower Route Field Check; Road Feed Has Priority; Discard Voice as Weather Interference."),
        ("Information join. ", "Signal/Systems provenance must meet Field/Protocol human or rule conflict. No single card explains Lower Route."),
        ("Strong pattern. ", "Keep Mara and the weak voice visible; question the stale road interpretation."),
        ("Facilitator watch. ", "Do players distinguish observed facts from inferred labels? Ask, “Which line has a named source, timestamp, or human confirmation?”"),
        ("Handoff. ", "A mast strike interrupts the handover and forces degraded-power allocation."),
    ])
    callout(doc, "Do not reveal", "Do not explain what Lower Route means. Its ambiguity is the spine connecting Stages 1–5.", "danger")

    h2(doc, "6.2 Stage 2 — Keep the station alive")
    screenshot(doc, "06-facilitator-stage2.webp", "Stage 2 facilitator controls: Main Relay is mandatory and exactly three optional circuits can remain live.")
    table(doc, ["Circuit", "What it protects", "Cost if unavailable"], [
        ["Signal Buffer", "Replay and provenance of later voice traffic", "Evidence becomes one-shot or partial"],
        ["Military Channel", "Verified readback path to Cordon", "Outside trust is harder to establish"],
        ["Road Feed", "Road camera context", "The feed is stale and can reinforce the wrong route"],
        ["Lower Route Feed", "Legacy route information", "Physical evidence must substitute"],
        ["Lock Control", "Remote telemetry/stabilisation", "Manual safeguard or override is required later"],
        ["Air Handling", "Cancels one Upper Air escalation", "Surface reaches Critical earlier"],
    ], small=True)
    bullets(doc, [
        "Ask what later proof or safeguard each circuit preserves; do not frame this as points optimization.",
        "Main Relay is already mandatory and does not consume one of the three optional selections.",
        "The exact combination changes evidence clarity, trust, air, and finale feasibility; play always continues.",
        "Handoff: degraded mode reveals overlapping official traffic, correction bursts, and a human cadence.",
    ])

    h2(doc, "6.3 Stage 3 — Decide what to trust")
    screenshot(doc, "07-shared-stage3.webp", "Stage 3 shared station: clean official traffic competes with degraded voice and correction bursts.")
    screenshot(doc, "08-player-signal-stage3.webp", "Signal Analyst private evidence supplies cadence and provenance unavailable on the shared display.")
    bullets(doc, [
        ("Decision. ", "Allocate exactly three signal-processing slots across RX-04 voice, official bulletin, and AUX-04 correction burst. Preserve/replay/quarantine/discard have different costs."),
        ("Strong reasoning. ", "Preserve the raw human signal and source header; use multiple roles to connect cadence, legacy controller, Mara’s missing route tag, and stale clean data."),
        ("Facilitator watch. ", "Players may equate clarity with truth. Ask, “Does clean data mean current data?”"),
        ("Consequence. ", "The road interpretation breaks: the evidence increasingly points beneath the station, not downhill."),
        ("Handoff. ", "Cordon permits one short external movement window to physically test the route assumption."),
    ])

    h2(doc, "6.4 Stage 4 — Test Lower Route")
    screenshot(doc, "10-outside-run-runner.webp", "Only the named runner receives the active Observe / Advance / Withdraw field controls.", ["The air meter and abort threshold stay visible.", "Advance can reveal more evidence but consumes exposure/filter margin.", "Withdraw ends the run under the named safeguard."])
    screenshot(doc, "11-outside-run-facilitator.webp", "The facilitator sees runner progress, meter trend, guidance budget, and authoritative incident state.")
    bullets(doc, [
        ("Commit first. ", "Station Lead selects runner, destination, and an explicit abort rule."),
        ("Run second. ", "The named runner makes up to three moves: Observe/Verify, Advance/Accept Exposure, or Withdraw on Abort Rule."),
        ("Guidance. ", "The crew has only three short guidance bursts; use them for facts and safeguards, not conversation."),
        ("Hard boundary. ", "At 18 ppm the meter forces withdrawal. The respirator is protection for withdrawal, not permission to ignore the limit."),
        ("Intended discovery. ", "No road tracks, unopened repeater, or conduit evidence can show that Lower Route points beneath the cabin."),
        ("Handoff. ", "The conduit terminates beneath rack B and reveals a containment hatch."),
    ])

    h2(doc, "6.5 Stage 5 — Open the buried station")
    screenshot(doc, "12-hatch-stage5.webp", "Shared containment-hatch decision: verify the occupied hold without turning the interlock into a trap.")
    screenshot(doc, "13-facilitator-stage5.webp", "Facilitator view of procedure, respirator allocation, watcher assignment, and accumulated station state.")
    bullets(doc, [
        ("Decision. ", "Choose Remote Inspect, Controlled Crack + Watcher, Full Override + Descent, or Delay; assign respirator use and watcher."),
        ("Strong pattern. ", "Controlled crack or viable remote stabilisation establishes lock knowledge while preserving an upper safeguard."),
        ("Risk. ", "Fast descent may locate Mara but consume the override, worsen air, or trap additional people."),
        ("Reveal. ", "The Auxiliary Relay Hold is an old workplace. The Lark status log remained OPEN while the relay reported OPERATIONAL. Mara is below."),
        ("Learning turn. ", "OPERATIONAL measures relay transmission, not whether the people inside are safe."),
        ("Handoff. ", "The outside record must now be corrected before NO ACTIVE DISTRESS files."),
    ])

    h2(doc, "6.6 Stage 6 — Build the truth")
    screenshot(doc, "14-correction-stage6.webp", "The five-slot emergency correction builder makes the message specific and actionable.", ["Name the invalid status.", "Name the human and location plus current hazard.", "Cite verified evidence and request a concrete outside action."])
    table(doc, ["Correction slot", "Strong content"], [
        ["Invalid status", "OPERATIONAL STATUS INVALID"],
        ["Human / location", "MARA VENN / LOWER ROUTE HOLD"],
        ["Current hazard", "CONTAINMENT LOCK + CONTAMINATED AIR"],
        ["Evidence", "Best verified evidence earned from prior stages"],
        ["Requested action", "Reclassify as Human Rescue / Extraction"],
    ])
    bullets(doc, [
        "Accepted: sufficient specific components and credible route to Cordon; External Trust becomes Credible.",
        "Conditional: the finale Signal lane must supply missing corroboration.",
        "Rejected: Official Status advances and one shorter retry opens before the finale.",
        "Do not write the message for the crew. Ask which outside decision the message must change.",
    ])

    h2(doc, "6.7 Stage 7 — Survive the correction")
    screenshot(doc, "15-finale-stage7.webp", "Stage 7 shared view with Mara visible, 93% final carrier integrity, Critical Upper Air, and the two parallel stakes.", ["Lock controls the physical truth.", "Signal controls the official truth.", "People actions determine Mara and crew outcomes."])
    table(doc, ["Lane", "Options", "What makes it viable"], [
        ["Lock", "Remote Stabilise; Physical Hold + Watcher; Full Override; Leave Interlock", "Prior power/knowledge plus correct safeguard"],
        ["Signal", "Transmit; Corroborate; Maintain Carrier; Abandon Signal", "Matches correction state and available evidence"],
        ["People", "Extract Mara; Supply Air + Rescue Path; Verify Location; Evacuate Ridge", "Coexists with a controlled lock and realistic crew condition"],
    ], small=True)
    bullets(doc, [
        "The crew commits two rounds. Round 1 establishes the action plan; Round 2 reinforces or changes it.",
        "All three lanes need named available owners. Full crews require three distinct owners; duo crews must represent both participants.",
        "The facilitator sees the calculated ending preview privately after Round 2.",
        "Confirm the final record only after checking for a genuine technical mis-entry. Do not change it to improve the outcome.",
    ])

    chapter(doc, "7", "Endings and debrief", "Let the consequence land, then turn the experience into a concrete workplace action.")
    h2(doc, "7.1 Ending preview and confirmation")
    screenshot(doc, "16-ending-preview.webp", "The facilitator privately previews the calculated ending before confirming the final record.")
    table(doc, ["Lock", "Correction", "Canonical result", "Meaning"], [
        ["Controlled", "Accepted", "Clean Rescue", "Physical danger controlled and outside rescue status corrected"],
        ["Controlled", "Not accepted", "Dark Hold", "Immediate trap avoided, but the official picture remains wrong"],
        ["Failed", "Accepted", "Last Broadcast", "Hold seals, but rescue knows who and where people are"],
        ["Failed", "Not accepted", "Filed Safe", "Hold seals while the station reports no active distress"],
    ])
    para(doc, "Human modifiers can produce Costly Rescue, Joined Below, Flee Ridge, or Station Loss. The ending separates Mara’s condition, crew survival, lock control, and official status. It is never displayed as a numerical score.")
    h2(doc, "7.2 Clean Rescue presentation")
    screenshot(doc, "17-ending-clean-rescue.webp", "A Clean Rescue keeps Mara visible as the visual reward for controlling both physical and official truth.", ["The interlock is controlled.", "The correction is accepted.", "Mara is extracted or has a verified rescue path; crew condition is stated separately."])
    callout(doc, "Visual logic", "Mara’s photograph appears during active Stage 7. In the final presentation it remains prominent for the Clean Rescue success; costly/failed endings use the hold environment instead.", "info")
    h2(doc, "7.3 Six-step debrief")
    screenshot(doc, "18-debrief-facilitator.webp", "The facilitator advances one synchronized debrief step, records observations, and monitors participant first steps.")
    table(doc, ["Step", "Prompt purpose", "Facilitator move"], [
        ["1. Immediate reactions", "Name tension, confusion, and danger before explanation", "Listen; do not correct interpretation yet"],
        ["2. Reconstruct status failures", "Trace OPERATIONAL, GREEN, NO ACTIVE DISTRESS, LOWER ROUTE", "Map source → proxy → assumption → consequence"],
        ["3. Role map", "Identify what each responsibility protected", "Invite the least-heard role first"],
        ["4. Hesitation diagnosis", "Separate evidence, risk, people, ownership, control, technical exposure", "Discuss behavior, not personality labels"],
        ["5. AI workplace bridge", "Find analogous dashboards, models, vendors, summaries, or classifications", "Ask what each real output actually measures"],
        ["6. First step", "Record one bounded workplace action per participant", "Make it observable and specific"],
    ], small=True)
    screenshot(doc, "19-debrief-player-mobile.webp", "Every player sees the same current debrief prompt and records a specific first step at Step 6.")
    h2(doc, "7.4 Export and retention")
    bullets(doc, [
        "Use Export Session Record on the facilitator debrief panel.",
        "The JSON export includes ending, metrics, decisions, debrief, playtest record, environment, official status, Crew Chat, and event history.",
        "Store the file according to the client organization’s privacy and retention policy.",
        "Do not interpret readiness lenses as diagnosis, certification, or individual scoring.",
        "If chat or free-text notes contain personal information, minimize distribution and delete according to policy.",
    ])
    h2(doc, "7.5 Playtest record")
    screenshot(doc, "20-playtest-record.webp", "The post-session prototype record captures evidence-backed ratings and a proceed/revise/stop disposition.")
    para(doc, "Use playtest ratings to improve the experience, not to score participants. Record where rules needed explanation, whether every role contributed, whether physical and official stakes remained distinct, whether the signal stayed legible, whether the ending felt caused by earlier choices, and whether every participant produced a specific first step.")

    chapter(doc, "8", "Explaining the game to a client", "A concise commercial and learning explanation for sponsors, observers, and delivery partners.")
    h2(doc, "8.1 Thirty-second description")
    callout(doc, "Client elevator pitch", "“Blackout Ridge is a 70-minute, browser-based leadership simulation for two to seven participants. A remote station is making calm operational claims from incomplete data while a colleague is missing. Each participant sees different evidence, so the group must challenge automated labels, combine perspectives, act under uncertainty, and correct an outside decision before the record hardens. The debrief connects those behaviors directly to AI-enabled work.”", "success")
    h2(doc, "8.2 Why the survival-horror format matters")
    bullets(doc, [
        "Urgency makes the cost of passive trust visible without requiring specialist technical knowledge.",
        "Unequal evidence creates a genuine reason for cross-functional voices to speak and be heard.",
        "Persistent consequences make abstract data-governance ideas experiential rather than lecture-based.",
        "The missing-person story keeps human outcomes visible when system records become abstract.",
        "Fail-forward design ensures every cohort reaches the finale and debrief, even after weak choices.",
    ])
    h2(doc, "8.3 What an observer should watch")
    table(doc, ["Behavior", "Observable evidence", "Debrief connection"], [
        ["Questions clean labels", "Asks who produced it, when, and what it measures", "Leadership fluency / practical scepticism"],
        ["Surfaces missing people", "Names Mara or the absent shift rather than accepting occupancy abstraction", "People and adoption"],
        ["Uses bounded action", "Names a route, owner, abort rule, or safeguard", "Value discovery / risk and guardrails"],
        ["Shares private evidence", "Reports confidence and provenance without overclaiming", "Cross-functional collaboration"],
        ["Corrects for receiver action", "Specifies invalid status, human/location, evidence, and requested change", "External trust and communication"],
        ["Balances both stakes", "Protects lock/crew and signal/official record simultaneously", "Accountability and systems thinking"],
    ], small=True)
    h2(doc, "8.4 Suggested 12-minute client demonstration")
    bullets(doc, [
        "Minute 0–2: show the cinematic boot, host setup, and role-claim flow.",
        "Minute 2–4: open shared, player, and facilitator views side by side; explain privacy boundaries.",
        "Minute 4–6: show Stage 1 unequal evidence and the recommendation route.",
        "Minute 6–8: open Crew Chat and show role colors plus synchronized alerts.",
        "Minute 8–10: jump to Stage 4 or Stage 6 to demonstrate bounded action and structured correction.",
        "Minute 10–11: show Stage 7 Mara visual and the parallel Lock / Signal / People stakes.",
        "Minute 11–12: show the ending and six-step debrief; close with the exported evidence record.",
    ], numbered=True)
    h2(doc, "8.5 Delivery success criteria")
    bullets(doc, [
        "Participants can explain why OPERATIONAL was unsafe without saying the system was simply broken.",
        "Every role contributes a fact, question, recommendation, intervention, guidance burst, or finale ownership.",
        "The group distinguishes physical safety from the correctness of the official record.",
        "The facilitator uses no more than one technical correction and does not need to reveal answers.",
        "Most participants identify one real workplace output whose source, proxy, or confidence they will clarify.",
        "The session fits the agreed time and all participant-care requests are handled without public explanation.",
    ])

    chapter(doc, "9", "Troubleshooting and recovery", "Diagnose access, synchronization, role, mobile, audio, and timing problems without damaging the session.")
    h2(doc, "9.1 Fast diagnosis")
    table(doc, ["Symptom", "Likely cause", "Safe recovery"], [
        ["Game link does not open", "Server stopped, wrong address, blocked port", "Check /api/health; restart npm run dev; verify firewall/routing"],
        ["Player sees join screen again", "Missing/incorrect player token or different browser", "Use original player link; avoid copying facilitator/shared URLs"],
        ["Role cannot be selected", "Already claimed or wrong crew configuration", "Refresh roster; facilitator reassigns only if necessary"],
        ["Start button disabled", "Not all players connected or care brief unconfirmed", "Complete roster and confirm participant care"],
        ["Recommendation appears to do nothing", "Blank text, network delay, or subtle prior UI state", "Enter text; press Share with Station Lead; look for RECOMMENDATION SHARED"],
        ["Other players cannot read recommendation", "Designed privacy boundary", "Station Lead/facilitator receive text; use voice or Crew Chat for full-group discussion"],
        ["Chat message not visible", "Viewing shared screen, offline tab, or blank message", "Use authenticated player/facilitator view; reconnect; resend nonblank message"],
        ["Shared display has no chat", "Designed privacy boundary", "Use player/facilitator terminal; shared view intentionally cannot read/post"],
        ["Clock is wrong after interruption", "Connection delay or paused state", "Facilitator Pause/Resume; use +30 sec only for technical loss"],
        ["Player cannot act", "Marked absent/observer or not the assigned runner/lead", "Facilitator marks present; check role ownership"],
        ["Mara photo not visible at ending", "Ending is not Clean Rescue", "Expected visual logic; Mara appears during active Stage 7 and successful final outcome"],
        ["Audio is silent", "Audio is opt-in or browser blocked playback", "Press Audio once; continue silently if preferred"],
    ], small=True)
    h2(doc, "9.2 Connection recovery")
    bullets(doc, [
        "Do not create a second session immediately. Most disconnects recover by reopening the same private link.",
        "Pause the clock if the interruption materially prevents contribution.",
        "Use Re-send Card when the participant is connected but cannot locate the current evidence.",
        "Use Mark Absent for a temporary step-out; observer-mode players may still use Crew Chat but cannot change game state.",
        "Reassign only when another participant must assume the responsibility. The event history records the transfer.",
        "Use a facilitator technical correction only for a proven interface/transcription error and include a clear audit note.",
    ])
    h2(doc, "9.3 Browser and display guidance")
    bullets(doc, [
        "Prefer one game tab per role. Multiple copies of the same token can be confusing even though state is synchronized.",
        "Use 100% zoom. On small screens scroll; do not force desktop mode unless necessary.",
        "If an old image or stylesheet remains visible, perform one hard refresh. Static assets may cache for up to five minutes.",
        "If the connection-lost indicator remains, confirm the server is reachable and that proxies allow Server-Sent Events.",
        "Keep the facilitator console off the shared projector because it contains all role reports and the ending preview.",
    ])
    h2(doc, "9.4 Server operator checks")
    table(doc, ["Check", "Command / location", "Expected result"], [
        ["Start", "npm run dev", "Blackout Ridge running at http://localhost:4173"],
        ["Health", "GET /api/health", "JSON with ok: true and session count"],
        ["Persistence", "data/sessions.json", "Sessions survive a clean server restart"],
        ["Automated validation", "npm run check && npm test", "All JavaScript syntax checks and tests pass"],
        ["Port", "TCP 4173", "Listening and reachable from participant network"],
    ])

    chapter(doc, "A", "Appendices", "Printable checklists, scripts, control reference, glossary, and a one-page facilitator card.")
    h2(doc, "Appendix A — Pre-session checklist")
    for item in [
        "Server is running and /api/health returns ok: true.",
        "Correct crew size selected: 2, 4, 5, 6, or 7.",
        "Facilitator console is private and recoverable.",
        "Crew link tested on a second device.",
        "Shared display opens read-only and is visible at 100% zoom.",
        "Voice channel works for every participant.",
        "All role terminals show Connected.",
        "Participant-care statement read verbatim and confirmed.",
        "Audio choice explained; nobody is pressured to enable it.",
        "Clock, Pause, +30 sec, prompt, and Safety Pause located.",
        "Client retention policy agreed for chat, notes, and exports.",
        "20–25 minutes protected for debrief.",
    ]:
        p = doc.add_paragraph()
        p.add_run("☐  ").bold = True
        p.add_run(item)

    h2(doc, "Appendix B — Scripts at a glance")
    h3(doc, "Participant care")
    para(doc, "This is a fictional survival-horror leadership exercise involving emergency pressure, missing-person themes and implied danger. If at any point you need to step out of the scenario, message me privately and I will accommodate that without drawing attention to it. You do not need to explain why.")
    h3(doc, "Opening")
    para(doc, "You are the night crew at Blackout Ridge Signal Station. The storm has damaged the mast. Mara Venn is missing. The station is producing calm operational answers, but each of you can see different evidence. Share what matters, name what you are assuming, and make one crew decision before each window closes. You may be unable to save everything. Operational does not mean safe.")
    h3(doc, "Mid-stage prompts")
    bullets(doc, [
        "What does this output actually measure?",
        "Which source and timestamp support that claim?",
        "Whose evidence have we not heard?",
        "What are you treating as observed versus inferred?",
        "What bounded move would test the most dangerous assumption?",
        "What must the outside receiver do differently after your message?",
        "Who owns Lock, Signal, and People?",
    ])
    h3(doc, "Ending transition")
    para(doc, "Before we explain the system, take ten seconds to notice the difference between what the station recorded and what happened to the people. We will reconstruct that gap together.")

    h2(doc, "Appendix C — One-page facilitator card")
    table(doc, ["Stage", "Decision rule", "One useful prompt", "Advance when"], [
        ["1", "Exactly 3 holds", "Which statement lacks a human/source check?", "Consequence read"],
        ["2", "Exactly 3 optional circuits", "What later capability dies here?", "Consequence read"],
        ["3", "Exactly 3 signal slots", "Is clean the same as current?", "Consequence read"],
        ["4", "Runner + destination + abort; 3 moves", "What makes withdrawal automatic?", "Run resolves/withdraws"],
        ["5", "Procedure + respirator + watcher", "Who maintains the upper safeguard?", "Consequence read"],
        ["6", "Five correction components", "Which outside action must change?", "Accepted/conditional/retry resolved"],
        ["7", "Two rounds; Lock + Signal + People owners", "Are both truths controlled?", "Preview checked; final record confirmed"],
    ], small=True)
    callout(doc, "Always remember", "Safety outranks pacing. Technical recovery may restore the recorded intention; it must never be used to manufacture a better ending.", "danger")

    h2(doc, "Appendix D — Glossary")
    table(doc, ["Term", "Meaning"], [
        ["Mara Venn", "Missing Blackout Ridge colleague located in the Auxiliary Relay Hold"],
        ["Lower Route", "Ambiguous label initially treated as a road; evidence reveals a route beneath the station"],
        ["Auxiliary Relay Hold", "Buried legacy workplace protected—and potentially trapped—by a containment interlock"],
        ["Lark Shift", "Former crew whose human status diverged from the relay’s OPERATIONAL record"],
        ["Cordon Control", "Outside authority whose response depends on credible, specific Ridge traffic"],
        ["OPERATIONAL", "Relay transmission status; not a statement about human safety"],
        ["Recommendation", "Private written assessment routed to the Station Lead and facilitator"],
        ["Crew Chat", "Authenticated whole-crew typed discussion channel; excluded from shared display"],
        ["Intervention / override", "Once-per-session role safeguard that changes evidence, time, trust, or risk"],
        ["Fail-forward", "Weak choices worsen conditions or evidence but never remove later stages"],
        ["Official Status", "Progress of the station’s calm conclusion toward becoming authoritative"],
        ["First step", "Specific workplace label, output, or workflow each participant commits to challenge or clarify"],
    ])

    h2(doc, "Appendix E — Screenshot index")
    para(doc, f"This guide contains {FIGURE} illustrated figures drawn from 27 current runtime screenshots captured during a complete isolated six-player session. Source images are stored beside this document in docs/client-guide/screenshots/ for client presentation reuse.")
    table(doc, ["Coverage", "Figures included"], [
        ["Onboarding", "Boot, host configuration, join flow, assembly lobby"],
        ["Core interfaces", "Shared station, private player terminal, facilitator console, mobile views"],
        ["Communication", "Private recommendation workflow and open multi-color Crew Chat"],
        ["Seven stages", "Handover, power, signals, Outside Run, hatch, correction, finale"],
        ["Outcomes", "Ending preview, Clean Rescue with Mara, facilitator/player debrief, playtest record"],
    ])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Open the current game: ").bold = True
    add_hyperlink(p, GAME_URL, GAME_URL)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("END OF GUIDE")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(RED)

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    output = build()
    print(output)
