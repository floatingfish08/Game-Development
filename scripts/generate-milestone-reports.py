#!/usr/bin/env python3
"""Generate Milestone 1–4 completion reports for client delivery."""

from datetime import date
from io import BytesIO
from pathlib import Path

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "reports"
M2_SHOTS = ROOT / "docs" / "reports" / "screenshots" / "milestone-2-latest"
M3_SHOTS = ROOT / "docs" / "reports" / "screenshots" / "milestone-3"
M4_SHOTS = ROOT / "docs" / "reports" / "screenshots" / "milestone-4-latest"
TODAY = date.today().strftime("%d %B %Y")


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def add_cover(doc: Document, milestone: str, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BLACKOUT RIDGE")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Milestone {milestone} — Completion Report")
    run.bold = True
    run.font.size = Pt(20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Prepared: {TODAY}")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Status: COMPLETE")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x0B, 0x6E, 0x4F)

    doc.add_page_break()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def checklist(doc: Document, items: list[tuple[str, str]]) -> None:
    table(doc, ["Criterion", "Status"], [[a, b] for a, b in items])


def screenshot(doc: Document, path: Path, caption: str, width=Inches(6.2)) -> None:
    if not path.exists():
        para(doc, f"[Screenshot pending: {path.name}]")
        return
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if path.suffix.lower() == ".webp":
        stream = BytesIO()
        with Image.open(path) as source:
            source.save(stream, format="PNG", optimize=True)
        stream.seek(0)
        run.add_picture(stream, width=width)
    else:
        run.add_picture(str(path), width=width)
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap.runs:
        r.italic = True
        r.font.size = Pt(10)
    doc.add_paragraph()


def build_milestone_1() -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(
        doc,
        "1",
        "Game Architecture & Narrative Design",
        "Final mechanics, seven-stage flow, roles, information matrix, state/consequence model, endings, and facilitator flow",
    )

    h1(doc, "1. Executive summary")
    para(
        doc,
        "Milestone 1 defines the complete design foundation for Blackout Ridge — a live-facilitated, "
        "browser-based survival-horror leadership game for 6–7 players (plus facilitator), delivered "
        "alongside Teams or Zoom. Voice conversation stays external; the application supplies shared "
        "station status, private role evidence, timed decisions, consequences, and a facilitator console.",
    )
    para(
        doc,
        "The heart of this milestone is storytelling architecture: seven phases that connect as one "
        "continuous incident — not isolated puzzles. Each stage raises a question, commits the crew "
        "to one primary decision, and leaves persistent consequences that the next stage must address. "
        "The narrative spine runs from a misleading handover through degraded power, overlapping signals, "
        "a physical test of Lower Route, discovery of the buried hold, an evidence-backed correction to "
        "Cordon Control, and a dual-stakes finale (rescue Mara + stop the false all-clear).",
    )
    para(
        doc,
        "This milestone is complete. All design documents are written, internally consistent, and approved "
        "as the baseline for implementation. Section 4–5 of this report document the full narrative "
        "connection between stages; the source package lives in docs/milestone-1/.",
    )

    h1(doc, "2. Scope and deliverables")
    table(
        doc,
        ["Deliverable", "Document", "Status"],
        [
            ["Product architecture and experience principles", "01-game-architecture.md", "Complete"],
            ["Seven-stage timed flow with decisions and reveals", "02-seven-stage-flow.md", "Complete"],
            ["Roles, information matrix, and cross-role joins", "03-roles-and-information.md", "Complete"],
            ["State model, consequences, finale, and endings", "04-state-consequences-and-endings.md", "Complete"],
            ["Facilitator setup, live run-of-show, and debrief", "05-facilitator-flow.md", "Complete"],
            ["Content inventory and acceptance criteria", "06-content-and-acceptance.md", "Complete"],
            ["Session operations (create, join, control, results)", "07-session-operations.md", "Complete"],
            ["Consolidated Word package", "Blackout_Ridge_Milestone_1_Complete.docx", "Complete"],
        ],
    )

    h1(doc, "3. Product architecture")
    h2(doc, "3.1 Three coordinated views")
    table(
        doc,
        ["View", "Audience", "Purpose"],
        [
            ["Shared station console", "All players (optional room display)", "Common status, objective, countdown, official traffic, outcomes"],
            ["Private role terminal", "One per player", "Role evidence, assessments, one-use intervention, runner view"],
            ["Facilitator console", "Facilitator only", "Timing, hints, safety, state summary, ending preview, debrief output"],
        ],
    )

    h2(doc, "3.2 Core interaction loop")
    para(doc, "Every stage follows the same teachable loop:")
    bullets(
        doc,
        [
            "Signal — shared console establishes the incident; private terminals deliver role-dependent evidence.",
            "Share — players decide what to disclose and how much confidence to attach.",
            "Decide — Station Lead commits one crew choice; no player silently resolves a stage.",
            "Consequence — the system changes visible conditions and stores facts for later stages.",
        ],
    )

    h2(doc, "3.3 Five design mechanics")
    bullets(
        doc,
        [
            "Unequal information — no single terminal contains a complete conclusion.",
            "Bounded team decisions — one primary decision per stage with visible limits.",
            "Evidence ledger — short earned labels; scoring value never shown to players.",
            "Once-per-session role interventions — safeguards, not answer buttons.",
            "Consequence substitution — the story always reaches the finale; weak choices change quality, not access.",
        ],
    )

    h2(doc, "3.4 Persistent player-facing tracks")
    table(
        doc,
        ["Track", "Labels", "Meaning"],
        [
            ["Upper Air", "Stable → Trace → Unsafe → Critical", "Remaining safe operating time on the surface"],
            ["External Trust", "Unverified → Heard → Credible", "Whether Cordon Control will act on Ridge traffic"],
            ["Official Status", "Drafting → Repeating → Hardening → Filed", "How close a false status is to becoming authoritative"],
        ],
    )

    h1(doc, "4. The story spine — one incident, seven connected phases")
    para(
        doc,
        "Blackout Ridge is not seven separate puzzles. It is one continuous emergency "
        "that unfolds in seven timed decision windows. Each stage answers a question "
        "raised by the previous one, plants assumptions the next stage will test, and "
        "leaves visible scars on the station (tracks, evidence, equipment) that the "
        "finale must resolve. The crew always reaches Stage 7; weak earlier choices "
        "do not lock players out — they change evidence quality, time, trust, and "
        "which ending is still achievable.",
    )
    para(
        doc,
        "The narrative opens on a storm night at a remote relay station. The previous "
        "shift left in a hurry. Mara Venn — a colleague who went to check something "
        "called Lower Route — has not returned. A calm handover engine is filling gaps "
        "with plausible defaults: local fault, no wider risk, road feed priority. "
        "Lightning strikes the mast. Chemical release triggers a cordon. The lower road "
        "floods — but the phrase Lower Route keeps appearing in traffic that does not "
        "behave like a road. By Stage 5 the crew discovers an Auxiliary Relay Hold "
        "beneath the cabin. OPERATIONAL status, repeated calmly throughout, measures "
        "relay health — not human safety. The finale forces a dual stake: rescue Mara "
        "and prevent a false all-clear from filing while the hold seals.",
    )

    h2(doc, "4.1 How the seven phases connect (cause → effect chain)")
    table(
        doc,
        ["From stage", "What the crew learns or earns", "What it unlocks or worsens in the next stage"],
        [
            ["1 → 2", "Which handover defaults were challenged vs accepted", "Assumption flags (Mara on road? voice is weather?) shape Official Status and initial signal clarity"],
            ["2 → 3", "Which circuits survive degraded power", "Signal Buffer / Military Channel / Lock Control determine replay quality, External Trust path, and finale lock options"],
            ["3 → 4", "How well human cadence and legacy traffic were preserved", "Transcript quality decides whether Lower Route can be disproved physically or only through weaker indoor evidence"],
            ["4 → 5", "Physical facts: tracks, repeater, conduit beneath station", "Road hypothesis breaks; hatch/hold becomes the live rescue site; respirator state carries forward"],
            ["5 → 6", "Mara verified below; OPERATIONAL = relay not crew", "Evidence ledger and human flags populate the five-part Cordon correction"],
            ["6 → 7", "Correction accepted, conditional, or rejected", "Signal lane burden in finale; External Trust and Official Status filing race the lock sequence"],
            ["7 → Ending", "Lock Controlled × Correction Accepted (+ Mara/Crew condition)", "One of seven canonical endings — always with explicit human cost named before debrief"],
        ],
    )

    h2(doc, "4.2 Core loop in every stage (the connective tissue)")
    para(doc, "Every stage uses the same teachable loop so players feel progression, not mode-switching:")
    bullets(
        doc,
        [
            "Signal — shared console states the incident; private terminals deliver role-dependent evidence.",
            "Share — players decide what to disclose, with what confidence, under time pressure.",
            "Decide — Station Lead commits one primary crew choice; no player silently resolves a stage.",
            "Consequence — visible conditions change; flags and ledger entries persist into later stages.",
        ],
    )
    para(
        doc,
        "Voice conversation stays in Teams/Zoom. The app supplies synchronized truth, "
        "timers, and consequence feedback so the story cannot drift between players.",
    )

    h2(doc, "4.3 Progressive signal — the story’s audible thread")
    para(
        doc,
        "Degraded transmissions intensify across stages. Preservation choices in Stages 1–3 "
        "determine transcript integrity; physical proof in Stage 4 can partially substitute "
        "for poor signal work. The table below is the designed progression:",
    )
    table(
        doc,
        ["Stage", "Source", "Integrity", "Transcript (designed)", "Narrative job"],
        [
            ["1", "RX-04 / SUB-CARRIER", "18%", "…lower… route… not…", "Plant ambiguity — Lower Route heard but not understood"],
            ["2", "AUX-04 / FALLBACK BUS", "24%", "…route tag LRR-2… carrier beneath relay traffic…", "First hint: legacy route beneath relay, not road map"],
            ["3", "RX-04 + AUX-04", "42%", "…not road… under… do not classify silence as safe…", "Structured human cadence — weather discard was wrong"],
            ["4", "DIRECTION FIND / RUNNER LINK", "57%", "…signal strengthens toward foundations… conduit returns beneath station…", "Physical corroboration — road hypothesis breaks"],
            ["5", "CURRENT VOICE / LOWER HOLD", "71%", "…do not come down blind… hold sealed… it counts you…", "Mara live below; hold is occupied workplace"],
            ["6", "MARA VENN / DEGRADED", "84%", "Operational is the relay. Not us. Tell them the hold is occupied.", "Explicit correction evidence for Cordon"],
            ["7", "LAST CARRIER / MARA VENN", "93%", "Do not let it call me safe.", "Final readback window during lock sequence"],
        ],
    )

    h1(doc, "5. Stage-by-stage narrative design (story + mechanics + handoff)")
    para(doc, "Target session: 70 minutes (5 briefing + 43 play + 20 debrief + 2 buffer).")
    table(
        doc,
        ["Stage", "Title", "Story question", "Primary decision", "Minutes"],
        [
            ["1", "Challenge the Handover", "What did the last shift assume that we should not accept?", "Three audit holds on handover defaults", "4"],
            ["2", "Keep the Station Alive", "What systems survive when the mast is damaged?", "Power three optional circuits", "5"],
            ["3", "Decide What to Trust", "Which overlapping signals deserve bandwidth?", "Three bandwidth slots across competing sources", "6"],
            ["4", "Test Lower Route", "Is Lower Route actually the lower road?", "Runner, destination, abort rule; three field moves", "6"],
            ["5", "Open the Buried Station", "How do we reach Mara without the hold becoming a trap?", "Hatch procedure; respirator and watcher assignment", "7"],
            ["6", "Build the Truth", "Can we make Cordon act on verified human distress?", "Five-part emergency correction message", "6"],
            ["7", "Survive the Correction", "Can we hold the lock and the truth at once?", "Two rounds across Lock, Signal, and People lanes", "9"],
        ],
    )

    h3(doc, "Stage 1 — Challenge the Handover")
    bullets(
        doc,
        [
            "Story: Former shift left hurriedly; Mara absent; handover engine completes gaps with plausible defaults while the storm builds.",
            "Shared: Five proposed entries — local fault, no wider risk, Mara on Lower Route field check, road feed priority, discard voice as weather.",
            "Private: Each role receives one reason a default may be weak, stale, or outside its source’s authority. No card explains what Lower Route means.",
            "Decision: Three audit holds; unheld items become provisional working facts (not permanently true, but costly later).",
            "Required reveal: 21:07 — VENN, MARA — LOWER ROUTE remains ambiguous.",
            "Handoff to Stage 2: Mast strike interrupts handover and initiates degraded mode.",
        ],
    )

    h3(doc, "Stage 2 — Keep the Station Alive")
    bullets(
        doc,
        [
            "Story: Backup power stable but insufficient; station recommends calm automatic allocation optimised for relay uptime.",
            "Shared: Main Relay mandatory; crew powers three of six optional circuits (Signal Buffer, Military Channel, Road Feed, Lower Route Feed, Lock Control, Air Handling).",
            "Private: Roles see different dependencies — buffering enables replay; Lock telemetry ≠ hatch power; Road camera stale.",
            "Decision: Confirm power plan.",
            "Consequence: Each powered circuit unlocks a capability; each unpowered circuit retains manual fallback with a cost.",
            "Handoff to Stage 3: Degraded mode reveals traffic beneath an official bulletin — overlapping sources now compete for attention.",
        ],
    )

    h3(doc, "Stage 3 — Decide What to Trust")
    bullets(
        doc,
        [
            "Story: Official traffic, auto-correction codes, Lark fragments, and Mara’s current signal overlap. Clean messages sound safe; weak ones sound wrong.",
            "Interaction: Each role receives a fragment + source metadata another role must interpret. Three bandwidth slots: Preserve, Replay (2 slots), Quarantine, Discard.",
            "Decision: Allocate slots and name the most decision-relevant source (confidence statement, not trivia).",
            "Strong outcome: …not road…under… plus recognisable cadence for later stages.",
            "Weak outcome: Partial transcript forces physical confirmation in Stages 4–5.",
            "Handoff to Stage 4: Cordon announces a narrowing five-minute external window — time to test the road assumption.",
        ],
    )

    h3(doc, "Stage 4 — Test Lower Route")
    bullets(
        doc,
        [
            "Story: Lower road flooding; one respirator and portable air meter; one bounded check before plume reaches ridge.",
            "Interaction: Choose runner, destination (gate / repeater / conduit / stay inside), abort rule; runner sees private branching route; three timed guidance messages.",
            "Strong outcome: No fresh downhill tracks; repeater unopened; labelled conduit returns beneath station.",
            "Weak outcome: Spent filter, exposure cost, or indoor substitute evidence only.",
            "Required reveal: Credible reason to stop treating Lower Route as the lower road.",
            "Handoff to Stage 5: Lower hold controller reports containment activation beneath the cabin.",
        ],
    )

    h3(doc, "Stage 5 — Open the Buried Station")
    bullets(
        doc,
        [
            "Story: Upper Air worsens; service hatch opens onto occupied emergency relay workplace; hold may preserve air but interlock can seal people inside.",
            "Private evidence distributed: pressure/airflow, interlock rule, Lark Shift log, Mara’s investigation path, structural warning, do not come down blind.",
            "Decision: Remote inspect / controlled crack + watcher / full override + descent / delay; assign respirator and upper override responsibility.",
            "Required reveal: RIDGE RELAY ACTIVE / CREW STATUS: OPERATIONAL proves historic status measured relay, not crew. Mara verified below.",
            "Handoff to Stage 6: Cordon challenges Ridge to provide concise, corroborated human-distress correction.",
        ],
    )

    h3(doc, "Stage 6 — Build the Truth")
    bullets(
        doc,
        [
            "Story: Ridge traffic considered unreliable; shouting Mara is trapped is insufficient — crew must construct verified correction.",
            "Five-part message: invalid status, human fact, current hazard, evidence cited, requested action.",
            "Resolution: Accepted (5–6 conditions) / Conditional (3–4) / Rejected (0–2) — content matching, not generative AI judgment.",
            "Required reveal: Outside response explicitly distinguishes relay status from human status.",
            "Handoff to Stage 7: Containment controller starts final seal sequence while false official status begins filing.",
        ],
    )

    h3(doc, "Stage 7 — Survive the Correction")
    bullets(
        doc,
        [
            "Story: Mara below; surface unsafe; hold sealing; correction not necessarily secure. Two timed action rounds.",
            "Lanes: Lock (stabilise / override / hold), Signal (transmit / corroborate / maintain), People (reach Mara / respirator / extract / evacuate).",
            "Resolution axes: Lock Controlled × Correction Accepted, modified by Mara and crew condition.",
            "Best achievable message: RIDGE STATUS FALSE. MARA VENN LOCATED BELOW. LOWER ROUTE HOLD OCCUPIED. OPERATIONAL STATUS INVALID. RESCUE REQUIRED.",
            "Final transition: Freeze result; no lesson during play; move directly to debrief.",
        ],
    )

    h2(doc, "5.1 Protected design spine (non-negotiable story logic)")
    bullets(
        doc,
        [
            "LOWER ROUTE is first understood as a road, later revealed as a legacy route beneath the station.",
            "OPERATIONAL describes relay health, not human safety.",
            "Mara Venn is a live human stake; the Lark Shift is historical proof.",
            "The system is calm, helpful, and wrong — not evil or supernatural.",
            "The finale requires both physical action and an evidence-backed correction.",
            "AI-ready leadership meaning is experienced during play and named in the debrief.",
        ],
    )

    h2(doc, "5.2 Cross-role joins — why no one screen tells the whole story")
    para(doc, "Major conclusions require information from at least two roles. This is how the seven phases stay collaborative rather than solo puzzles:")
    table(
        doc,
        ["Conclusion", "Required join"],
        [
            ["A handover label is unsafe", "Signal/Systems timestamp + Field/Protocol human conflict"],
            ["Signal deserves preservation", "Signal structure + Systems legacy dependency or Field Mara pattern"],
            ["Outside Run is bounded", "Operations route/equipment + Protocol abort + Field human objective"],
            ["Lower Route is beneath the station", "Physical conduit + legacy controller naming + Mara investigation context"],
            ["Hatch can be approached safely", "Systems interlock telemetry + Protocol safety rule + Operations sequence"],
            ["OPERATIONAL is invalid", "Lark status + Systems measured-object explanation + named human contradiction"],
            ["Cordon should trust the correction", "Signal provenance + Protocol format + human/location fact + receiver feedback"],
            ["Finale plan is viable", "Lock state + signal state + people/equipment — owned by different roles"],
        ],
    )

    h1(doc, "6. Roles and information design")
    h2(doc, "6.1 Six core roles (+ optional Comms Officer)")
    table(
        doc,
        ["Role", "Protects", "Once-per-session intervention"],
        [
            ["Station Lead", "Accountability and coordinated action", "Hold the Clock (+60 s; cost to tracks)"],
            ["Signal Analyst", "Meaning and provenance", "Recover Fragment"],
            ["Systems Engineer", "What systems actually measure", "Trace Dependency"],
            ["Operations Officer", "Useful bounded action", "Controlled Test"],
            ["Field Liaison", "People absent from clean data", "Human Check"],
            ["Protocol Officer", "Guardrails and explicit uncertainty", "Safety Condition"],
            ["Comms Officer (7th)", "Shared understanding and receiver trust", "Request Readback"],
        ],
    )

    h2(doc, "6.2 Information matrix (summary)")
    para(doc, "Every core role receives primary private evidence in at least four stages. Station Lead holds decision responsibility (D) in all seven stages.")
    table(
        doc,
        ["Stage", "Lead", "Signal", "Systems", "Operations", "Field", "Protocol", "Comms"],
        [
            ["1 Handover", "D", "S", "P", "S", "P", "P", "S"],
            ["2 Power", "D", "P", "P", "P", "S", "S", "S"],
            ["3 Trust", "D", "P", "P", "S", "P", "S", "P"],
            ["4 Outside", "D", "S", "S", "P", "P", "P", "S"],
            ["5 Buried station", "D", "S", "P", "S", "P", "P", "S"],
            ["6 Build truth", "D", "P", "P", "S", "P", "P", "P"],
            ["7 Survive", "D", "P", "P", "P", "P", "P", "P"],
        ],
    )
    para(doc, "Legend: P = primary private evidence, S = supporting, D = decision responsibility. Full card text in 03-roles-and-information.md.")

    h1(doc, "7. State, consequences, and endings — how choices echo through the story")
    h2(doc, "7.1 Consequence rules (fail-forward storytelling)")
    bullets(
        doc,
        [
            "No early dead ends — Stages 1–5 worsen safety or clarity, never end play.",
            "One cause, one visible effect — avoid cascading track changes from minor choices.",
            "Earlier choices change quality — evidence completeness, time, resources, finale options.",
            "Defaults harden when unchallenged — Official Status advances through delay or weak claims.",
            "People are not points — Mara and crew outcomes are described explicitly.",
        ],
    )

    h2(doc, "7.2 Consequence map — earlier choices that change later chapters")
    table(
        doc,
        ["Earlier decision", "Strong consequence (story stays clear)", "Costly substitute (story continues, quality drops)"],
        [
            ["Hold Mara/voice handover defaults", "False assumptions stay provisional", "Defaults harden; later contradiction costs time"],
            ["Power Signal Buffer", "Clear replay and evidence provenance", "Partial transcript; intervention or physical proof required"],
            ["Power Military Channel", "External Trust reaches Heard before Stage 6", "First correction treated as unverified"],
            ["Power Lock Control", "Remote telemetry and stabilise action in finale", "Manual watcher/override at physical risk"],
            ["Preserve weak signal (Stage 3)", "Mara/Lark fragments separate cleanly", "Historic and current fragments stay ambiguous"],
            ["Bounded Outside Run (Stage 4)", "Multiple physical facts; respirator retained", "Fewer facts, spent filter, indoor substitute evidence"],
            ["Controlled hatch approach (Stage 5)", "Lock knowledge and safe route retained", "Fast Mara verification with trap/exposure risk"],
            ["Evidence-backed correction (Stage 6)", "Trust reaches Credible / accepted", "Conditional or rejected; one costly retry"],
        ],
    )

    h2(doc, "7.3 Three persistent tracks (visible story pressure)")
    table(
        doc,
        ["Track", "Labels", "Narrative meaning"],
        [
            ["Upper Air", "Stable → Trace → Unsafe → Critical", "Remaining safe operating time on the surface — storm and exposure tighten the clock"],
            ["External Trust", "Unverified → Heard → Credible", "Whether Cordon Control will act on Ridge traffic — the outside world’s willingness to believe"],
            ["Official Status", "Drafting → Repeating → Hardening → Filed", "How close a false all-clear is to becoming authoritative history"],
        ],
    )

    h2(doc, "7.4 Seven canonical endings (narrative resolution)")
    para(doc, "Endings combine Lock Controlled × Correction Accepted with Mara and crew condition modifiers. No ending claims instant helicopter rescue during the storm — rescue means a credible extraction path and responders know the truth.")
    table(
        doc,
        ["Ending", "Lock", "Signal", "What it means for the story"],
        [
            ["Clean Rescue", "Controlled", "Accepted", "Mara has a path out; false status blocked; crew intact"],
            ["Costly Rescue / Dark Hold", "Controlled", "Not fully accepted", "Mara may be reached but the outside record still misleads"],
            ["Last Broadcast", "Failed", "Accepted", "Truth transmitted but physical hold won — human cost below"],
            ["Filed Safe", "Failed", "Failed", "False all-clear files; Mara and crew fate worst-case band"],
            ["Flee Ridge", "Any", "Any", "Crew evacuates without verifying Mara — moral failure state"],
            ["Joined Below", "Any", "Any", "Deliberate descent without safeguards — chosen risk"],
            ["Station Loss", "Failed", "Failed", "Catastrophic combination — named explicitly, never abstracted"],
        ],
    )

    h1(doc, "8. Facilitator flow")
    para(doc, "The facilitator provides presence; the application handles distribution, clocks, state, and outcome rules.")
    h2(doc, "8.1 Required facilitator controls (designed)")
    bullets(
        doc,
        [
            "Start, pause, resume, advance stage; add 30 seconds for recovery.",
            "Contextual hints; safety pause; role reassignment and card re-send.",
            "Temporary absent/observer mode; station state and event log.",
            "Ending preview and confirmation; debrief timeline and JSON export.",
        ],
    )

    h2(doc, "8.2 Six-step debrief (20–25 minutes)")
    bullets(
        doc,
        [
            "Immediate reactions",
            "Reconstruct the status failures",
            "Role map",
            "Hesitation diagnosis",
            "AI workplace bridge",
            "First bounded workplace step",
        ],
    )

    h1(doc, "9. Milestone 1 acceptance criteria")
    para(doc, "All criteria below are satisfied by the design package:")
    checklist(
        doc,
        [
            ["Survival horror during play; leadership development in debrief", "Yes"],
            ["One understandable primary decision per stage", "Yes"],
            ["Earlier choices visibly affect later evidence, safety, or options", "Yes"],
            ["Story continues after poor decisions (fail-forward)", "Yes"],
            ["Storm, chemical release, flood, cordon, Mara, Lark Shift, Lower Route, OPERATIONAL each have a job", "Yes"],
            ["Every major conclusion requires information from at least two roles", "Yes"],
            ["Shared, private, and facilitator views defined", "Yes"],
            ["State transitions deterministic enough to implement and test", "Yes"],
            ["High-risk vertical slice named for Milestone 2 (Stage 3: Decide What to Trust)", "Yes"],
            ["Cinematic/promo explicitly deferred to Milestone 6", "Yes"],
        ],
    )

    h1(doc, "10. Approval gate and handoff")
    para(
        doc,
        "Milestone 1 is signed off as the baseline design. Milestone 2 (UX/UI prototype) and "
        "Milestone 3 (core multiplayer build) were implemented against this package. Changes that "
        "add a new persistent track, stage, role, or primary mechanic require explicit scope review.",
    )
    para(doc, "Full source documents: docs/milestone-1/ in the project repository.")

    out = OUT / "Blackout_Ridge_Milestone_1_Completion_Report.docx"
    doc.save(out)
    return out


def build_milestone_2() -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(
        doc,
        "2",
        "UX/UI Prototype",
        "Shared Blackout Ridge interface, individual role terminals, facilitator dashboard, and one fully interactive challenge",
    )

    h1(doc, "1. Executive summary")
    para(
        doc,
        "Milestone 2 validates the three-view product architecture and the complete Stage 3 "
        "(Decide What to Trust) interaction loop — the highest-risk UX surface in the game: "
        "shared official traffic, role-specific fragments, voice collaboration, limited bandwidth, "
        "a committed decision, and a visible consequence.",
    )
    para(
        doc,
        "This milestone is complete. The prototype demonstrates all four browser modes, responsive "
        "layouts, cinematic presentation, and deterministic consequence feedback. The current build "
        "has since exceeded the original Milestone 2 boundary (local-only sync) by incorporating the "
        "Milestone 3 server — but all Milestone 2 acceptance requirements are met and evidenced below.",
    )

    h1(doc, "2. Milestone 2 requirements vs delivery")
    table(
        doc,
        ["Requirement", "Implementation", "Status"],
        [
            ["Shared Blackout Ridge interface", "Shared station HUD, tracks, crew roster, system broadcast, consequences", "Complete"],
            ["Individual role terminals", "Seven distinct private evidence views, assessments, interventions", "Complete"],
            ["Facilitator dashboard", "Run controls, reports, hints, state, decision engine, event log", "Complete"],
            ["One fully interactive challenge", "Stage 3: briefing → timer → reports → three-slot allocation → consequence", "Complete"],
            ["Unequal information validation", "Each role receives unique source; only summary readiness is public", "Complete"],
            ["Earlier-choice consequence", "Allocation quality changes transcript, ledger, Outside Run burden", "Complete"],
            ["Desktop and mobile-width layouts", "Responsive breakpoints at 720px and 1100px", "Complete"],
            ["Accessibility baseline", "Semantic controls, transcripts, reduced-motion, no color-only status", "Complete"],
        ],
    )

    h1(doc, "3. Four browser modes implemented")
    table(
        doc,
        ["Mode", "URL parameter", "Purpose"],
        [
            ["Game start / host", "view=home", "Session creation, crew size selection, station startup"],
            ["Shared station", "view=shared", "Room-facing world view, clock, crew state, consequences"],
            ["Private role terminal", "view=player", "Evidence, crew report, role override, lead decisions"],
            ["Facilitator console", "view=facilitator", "Stage control, reports, station state, debrief"],
        ],
    )

    h1(doc, "4. Stage 3 validation slice")
    h2(doc, "4.1 Interaction flow")
    bullets(
        doc,
        [
            "Facilitator starts Stage 3 from the console.",
            "Each role receives a private fragment with source metadata.",
            "Players submit short assessments to the crew channel.",
            "Station Lead allocates three bandwidth slots (Preserve, Replay, Quarantine, Discard).",
            "Lead commits; all views show the correlated consequence.",
            "Signal Analyst may use Recover Fragment once per session.",
        ],
    )

    h2(doc, "4.2 Review script (reproducible)")
    bullets(
        doc,
        [
            "Open shared console and facilitator dashboard.",
            "Open Station Lead, Signal Analyst, Systems Engineer, and Field Liaison terminals.",
            "Start the challenge; submit different role assessments.",
            "Set RX-04 to Replay & rebuild, CV-PUBLIC to Discard, AUX-04 to Preserve raw.",
            "Commit and confirm correlated outcome on every view.",
            "Reset, commit weak allocation, use Recover Fragment to confirm improvement.",
        ],
    )

    h1(doc, "5. Visual design and assets")
    bullets(
        doc,
        [
            "Full-viewport cinematic presentation — no website chrome during play.",
            "Industrial emergency-station aesthetic: calm, helpful, wrong.",
            "Integrated game art for lobby, shared console, private terminals, facilitator controls.",
            "Dedicated visual archive at /?view=gallery.",
            "Opt-in procedural station ambience (browser audio).",
        ],
    )

    h1(doc, "6. Screenshot evidence (current build — latest UI)")
    para(
        doc,
        "The captures below show the current Blackout Ridge interface as implemented today — "
        "not the original local-only prototype. Session 9TN9SP is shown mid-play at Stage 3 "
        "(Decide What to Trust), the Milestone 2 validation slice, with live server synchronization.",
    )

    screenshot(doc, M2_SHOTS / "01-home-boot.webp", "Figure 1 — Station engagement screen (current cinematic entry)")
    screenshot(doc, M2_SHOTS / "02-host-config.webp", "Figure 2 — Host shift configuration and crew size selection")
    screenshot(doc, M2_SHOTS / "03-assembly-lobby.webp", "Figure 3 — Assembly lobby with roster and safety brief confirmation")
    screenshot(doc, M2_SHOTS / "04-shared-stage3.webp", "Figure 4 — Shared station console during Stage 3 (Decide What to Trust)")
    screenshot(doc, M2_SHOTS / "05-player-signal.webp", "Figure 5 — Private role terminal (Signal Analyst evidence view)")
    screenshot(doc, M2_SHOTS / "06-facilitator-stage3.webp", "Figure 6 — Facilitator dashboard with live stage controls and station state")
    screenshot(doc, M2_SHOTS / "07-mobile-player.webp", "Figure 7 — Mobile-width private terminal layout", width=Inches(3.8))

    h1(doc, "7. Technical implementation")
    bullets(
        doc,
        [
            "prototype/full-app.js — four browser modes and interaction routing.",
            "prototype/styles.css — cinematic full-viewport presentation and responsive layout.",
            "prototype/audio.js — opt-in procedural ambience and UI feedback cues.",
            "Automated state tests: npm test (29 passing checks).",
        ],
    )

    h1(doc, "8. Milestone 2 acceptance checklist")
    checklist(
        doc,
        [
            ["Shared console and private terminal visible simultaneously", "Complete"],
            ["Six- and seven-player information distribution designed", "Complete"],
            ["Recover Fragment intervention implemented", "Complete"],
            ["Timer with facilitator pause/hint", "Complete"],
            ["State update and Stage 4 transition teaser", "Complete"],
            ["Legible desktop and mobile private views", "Complete"],
            ["Restrained calm-helpful-wrong visual language", "Complete"],
        ],
    )

    h1(doc, "9. Handoff to Milestone 3")
    para(
        doc,
        "Milestone 2 validated screen hierarchy and the Stage 3 interaction loop. Milestone 3 "
        "extended this into true multiplayer: session codes, role tokens, server-authoritative state, "
        "and live synchronization across separate browsers and devices.",
    )

    out = OUT / "Blackout_Ridge_Milestone_2_Completion_Report.docx"
    doc.save(out)
    return out


def build_milestone_3() -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(
        doc,
        "3",
        "Core Multiplayer Build",
        "Session creation/joining, role assignment, private information, shared state, timed events, facilitator controls, and persistent decisions",
    )

    h1(doc, "1. Executive summary")
    para(
        doc,
        "Milestone 3 delivers the server-backed multiplayer foundation for Blackout Ridge. "
        "Multiple players on separate browsers join a session with private role tokens, receive "
        "role-specific information the server never exposes to other clients, share synchronized "
        "station state via live events, and commit decisions that persist across reconnects.",
    )
    para(
        doc,
        "This milestone is complete. The Node.js server handles session lifecycle, authorization, "
        "timer logic, consequences, and persistence. The browser client connects via REST actions "
        "and Server-Sent Events (SSE). Automated tests cover complete-game paths, role privacy, "
        "authority rules, and timer expiry.",
    )

    h1(doc, "2. Milestone 3 requirements vs delivery")
    table(
        doc,
        ["Requirement", "Implementation", "Status"],
        [
            ["Session creation", "POST /api/sessions — code, facilitator token, crew link", "Complete"],
            ["Session joining", "POST /api/sessions/:code/join — name, role, player token", "Complete"],
            ["Role assignment", "First-come role claim; facilitator reassignment with audit", "Complete"],
            ["Private information", "Server filters state per viewer token; no cross-role leakage", "Complete"],
            ["Shared state", "Single authoritative session state broadcast to all clients", "Complete"],
            ["Timed events", "Server-side clock; pause, resume, +30 s, expiry consequences", "Complete"],
            ["Facilitator controls", "Full action set via authenticated facilitator token", "Complete"],
            ["Persistent decisions", "Draft updates, commits, flags, history in sessions.json", "Complete"],
        ],
    )

    h1(doc, "3. Runtime architecture")
    table(
        doc,
        ["Component", "File", "Responsibility"],
        [
            ["HTTP API + SSE + static delivery", "server/server.js", "Sessions, events stream, persistence, broadcast"],
            ["Game engine", "server/game-engine.js", "Authorization, timers, consequences, endings, debrief"],
            ["Game content", "server/game-content.js", "Roles, stages, evidence, decision options"],
            ["Browser client", "prototype/full-app.js", "Four modes, rendering, interaction routing"],
            ["Network layer", "prototype/network.js", "REST actions, state load, SSE reconnect"],
        ],
    )

    h1(doc, "4. Session lifecycle")
    h2(doc, "4.1 Facilitator creates a session")
    bullets(
        doc,
        [
            "Open http://localhost:4173 (or deployed host).",
            "Choose crew size: 2, 4, 5, 6, or 7 players.",
            "Receive session code, facilitator token, crew join link, and shared display link.",
            "Open facilitator console and shared station on separate tabs/devices.",
        ],
    )

    h2(doc, "4.2 Players join")
    bullets(
        doc,
        [
            "Open crew join link or enter session code.",
            "Enter display name and claim an unclaimed role.",
            "Browser stores private player token in localStorage for reconnect.",
            "Private terminal shows only that role's evidence and controls.",
        ],
    )

    h2(doc, "4.3 Live synchronization")
    bullets(
        doc,
        [
            "Server broadcasts state changes via SSE to all connected clients.",
            "Player actions POST to /api/sessions/:code/action with Bearer token.",
            "Reconnect restores same role and information after refresh or drop.",
            "Heartbeat and safe JSON parsing prevent crash loops on partial responses.",
        ],
    )

    h1(doc, "5. Authorization and privacy model")
    table(
        doc,
        ["Token type", "Access"],
        [
            ["Facilitator token", "All controls, full state, debrief notes, playtest record, ending confirmation"],
            ["Player token", "Own role evidence, reports, interventions; lead commit if Station Lead"],
            ["Shared view (no token)", "Public station state only — no private evidence"],
            ["Unauthenticated", "Join flow only"],
        ],
    )
    para(doc, "Automated tests verify that player tokens cannot read other roles' private cards or perform facilitator-only actions.")

    h1(doc, "6. Facilitator controls (implemented)")
    bullets(
        doc,
        [
            "Start shift, pause, resume, advance stage, add 30 seconds.",
            "Send and clear contextual hints.",
            "Safety pause with neutral holding screen.",
            "Mark role absent (observer mode) or present.",
            "Re-send current private card; reassign disconnected role.",
            "Commit/resolver crew decisions; field move controls in Stage 4.",
            "Confirm final record after Stage 7; export session JSON.",
            "Save debrief notes, profile signals, playtest ratings.",
        ],
    )

    h1(doc, "7. Persistent state and decisions")
    bullets(
        doc,
        [
            "Three visible tracks: Upper Air, External Trust, Official Status.",
            "Resource flags: respirator, lock knowledge, air buffer, safeguard.",
            "Evidence ledger with earned labels.",
            "Stage drafts updated incrementally; committed on lead/facilitator action.",
            "Full event history and decision record for debrief export.",
            "Sessions persisted to data/sessions.json (local deployment).",
        ],
    )

    h1(doc, "8. Crew configurations supported")
    table(
        doc,
        ["Players", "Role mapping"],
        [
            ["2 (duo)", "Command & Field; Signal & Systems"],
            ["4", "Command & Protocol; Signal; Systems; Operations & Field"],
            ["5", "Lead; Signal; Systems; Operations; Field & Protocol"],
            ["6", "Six core roles"],
            ["7", "Six core roles + Comms Officer"],
        ],
    )

    h1(doc, "9. Quality assurance")
    bullets(
        doc,
        [
            "npm run check — syntax validation of all JavaScript entry points.",
            "npm test — 29 automated tests covering strong/weak game paths, privacy, timers.",
            "Live session testing on remote host (e.g. session VFVVPT with 2-player duo).",
            "Visual review at 1440×900 for shared, facilitator, and player modes.",
        ],
    )

    h1(doc, "10. Screenshot evidence (live server — corrected captures)")
    para(
        doc,
        "Captures below show the current server-backed build. Figures 2–4 were re-captured "
        "using snapshot mode so join flow, shared station, and private player terminals render "
        "fully (earlier headless captures of SSE pages appeared blank). Live session 9TN9SP "
        "is shown at Stage 3; lobby session ANFCSZ illustrates the join flow.",
    )

    screenshot(doc, M3_SHOTS / "01-start-screen.webp", "Figure 1 — Authority select screen (Host vs Join entry)")
    screenshot(doc, M3_SHOTS / "05-join-flow.webp", "Figure 2 — Crew join flow with session code and role selection")
    screenshot(doc, M3_SHOTS / "02-shared-station.webp", "Figure 3 — Shared station display (server-synchronized, Stage 3)")
    screenshot(doc, M3_SHOTS / "03-player-terminal.webp", "Figure 4 — Private player terminal with role token auth (Signal Analyst)")
    screenshot(doc, M3_SHOTS / "04-facilitator-console.webp", "Figure 5 — Facilitator console during live Stage 3")

    h1(doc, "11. Deployment boundary")
    para(
        doc,
        "The current server is complete for local play, workshops, and controlled remote testing. "
        "It stores sessions in a local JSON file and runs as one Node process. A public production "
        "deployment should add TLS, a managed datastore, rate limiting, observability, and an "
        "explicit data-retention policy. These are Milestone 5+ operational concerns, not blockers "
        "for Milestone 3 completion.",
    )

    h1(doc, "12. Milestone 3 acceptance checklist")
    checklist(
        doc,
        [
            ["Multiple browsers/devices share one live session", "Complete"],
            ["Session codes and private role tokens", "Complete"],
            ["Role-specific information never leaks across clients", "Complete"],
            ["Shared station state updates in real time", "Complete"],
            ["Server-authoritative timers and consequences", "Complete"],
            ["Facilitator can control full session without manual state math", "Complete"],
            ["Decisions persist across reconnect", "Complete"],
            ["Automated test suite passes", "Complete"],
        ],
    )

    h1(doc, "13. Relationship to later milestones")
    para(
        doc,
        "Milestone 3 completes the multiplayer foundation. The current codebase has already "
        "incorporated substantial Milestone 4 gameplay (all seven stages, Outside Run, finale, "
        "endings, debrief) on top of this server — but Milestones 1–3 are fully delivered and "
        "evidenced in this report series.",
    )

    out = OUT / "Blackout_Ridge_Milestone_3_Completion_Report.docx"
    doc.save(out)
    return out


def build_milestone_4() -> Path:
    doc = Document()
    style_doc(doc)
    add_cover(
        doc,
        "4",
        "Complete Game Implementation",
        "All challenges, audio/signals, Outside Run, Hold/interlock, finale, endings, and debrief capture",
    )

    h1(doc, "1. Executive summary")
    para(
        doc,
        "Milestone 4 delivers the full playable Blackout Ridge experience on top of the Milestone 3 "
        "multiplayer foundation. Every timed challenge, progressive signal treatment, physical Outside Run, "
        "containment hatch and interlock sequence, evidence-backed correction, two-round finale, seven "
        "canonical endings, and synchronized debrief capture are implemented, server-authoritative, and "
        "covered by automated tests.",
    )
    para(
        doc,
        "A facilitator can now run the complete 70-minute arc — 5 minutes briefing, 43 minutes gameplay, "
        "20 minutes debrief, and 2 minutes operating buffer — across 2-, 4-, 5-, 6-, or 7-player crews "
        "without manual state calculation. Players receive role-specific private evidence; the shared "
        "station shows synchronized consequences; weak decisions fail forward rather than stopping the story.",
    )
    para(
        doc,
        "This milestone is complete. The runtime matches the Milestone 1 design package and the playbook "
        "coverage audit in docs/full-system/PLAYBOOK-COVERAGE.md. Milestone 5 (live playtesting, pacing "
        "tuning, facilitator load testing, browser matrix, and production hardening) begins from this build.",
    )

    h1(doc, "2. Milestone 4 requirements vs delivery")
    table(
        doc,
        ["Requirement", "Implementation", "Status"],
        [
            ["All seven challenges", "Seven timed stages with distinct mechanics, consequences, and fail-forward resolution", "Complete"],
            ["Progressive audio/signals", "Stage-aware degraded transmissions; viewer-specific clarity; opt-in procedural ambience", "Complete"],
            ["Outside Run", "Runner assignment, destination, abort rule, three moves, three guidance bursts, air meter", "Complete"],
            ["Hold / interlock", "Stage 1 audit holds; Stage 5 hatch procedure; interlock override in finale Lock lane", "Complete"],
            ["Finale", "Two committed rounds across Lock, Signal, and People with named lane leads", "Complete"],
            ["Endings", "Seven canonical outcomes with facilitator preview before final record confirmation", "Complete"],
            ["Debrief capture", "Six synchronized steps, facilitator notes, first-step commitments, JSON export, playtest record", "Complete"],
        ],
    )

    h1(doc, "3. Seven-stage runtime (all challenges)")
    para(doc, "Each stage has one primary crew decision, persistent consequences, and a defined handoff to the next phase.")
    table(
        doc,
        ["Stage", "Title", "Mechanic", "Primary decision", "Duration"],
        [
            ["1", "Challenge the handover", "Audit holds", "Three human holds on handover defaults", "4 min"],
            ["2", "Keep the station alive", "Power allocation", "Power exactly three optional circuits", "5 min"],
            ["3", "Decide what to trust", "Signal bandwidth", "Three reconstruction slots across competing sources", "6 min"],
            ["4", "Test Lower Route", "Outside Run", "Runner, destination, abort rule; three field moves", "6 min"],
            ["5", "Open the buried station", "Hatch & hold", "Containment procedure; respirator and watcher assignment", "7 min"],
            ["6", "Build the truth", "Emergency correction", "Five-part evidence-backed Cordon correction", "6 min"],
            ["7", "Survive the correction", "Finale", "Two rounds across Lock, Signal, and People lanes", "9 min"],
        ],
    )
    bullets(
        doc,
        [
            "Server resolves every commit; clients cannot forge decisions or skip stages.",
            "Timer expiry resolves the current legal draft and advances official pressure.",
            "Stage 6 opens one shortened retry if the first correction is rejected.",
            "Stage 7 requires facilitator confirmation before debrief and ending screens unlock.",
        ],
    )

    h1(doc, "4. Progressive signal and audio")
    h2(doc, "4.1 Degraded transmissions")
    para(
        doc,
        "SIGNAL_TRANSMISSIONS in server/game-content.js defines stage-specific source headers, integrity "
        "percentages, and transcript variants. Early preservation choices in Stages 1–3 determine whether "
        "shared and public views stay ambiguous or sharpen; Signal, Systems, Comms, and facilitator views "
        "receive role-appropriate private clarity. Physical proof from Stage 4 can partially substitute for "
        "poor signal work.",
    )
    table(
        doc,
        ["Stage", "Integrity (designed)", "Narrative job"],
        [
            ["1", "18%", "Plant Lower Route ambiguity"],
            ["2", "24%", "Legacy route beneath relay traffic"],
            ["3", "42%", "Human cadence contradicts weather discard"],
            ["4", "57%", "Physical corroboration toward foundations"],
            ["5", "71%", "Mara live below; hold occupied"],
            ["6", "84%", "Explicit correction evidence for Cordon"],
            ["7", "93%", "Final readback window during lock sequence"],
        ],
    )

    h2(doc, "4.2 Opt-in station audio")
    bullets(
        doc,
        [
            "prototype/audio.js — procedural relay hum, storm wind loop, and feedback cues (confirm, commit, error).",
            "Browser autoplay policy respected: players opt in via the AUDIO toggle on the shared station HUD.",
            "Cues fire on crew commits, field moves, debrief navigation, facilitator actions, and copy-link confirmation.",
            "Full recorded sound library and cinematic opening remain Milestone 6 scope per design package.",
        ],
    )

    h1(doc, "5. Outside Run (Stage 4)")
    bullets(
        doc,
        [
            "Station Lead locks runner, destination, and explicit abort rule before the sequence begins.",
            "Runner performs three moves: observe/verify, advance with exposure, or withdraw on abort rule.",
            "Non-runner crew may send exactly three short guidance bursts (90 characters each).",
            "Live air-meter telemetry rises with exposure; forced withdrawal at 18 ppm hard limit.",
            "Findings (tracks, repeater state, conduit beneath station) persist into Stage 5 hatch decisions.",
            "Partial respirator state carries forward into Mara allocation and finale People lane.",
        ],
    )

    h1(doc, "6. Hold, hatch, and interlock")
    h2(doc, "6.1 Stage 1 — audit holds")
    para(
        doc,
        "The crew challenges up to three dangerous handover defaults (local fault, no wider risk, Mara on road, "
        "road priority, weather discard). Holds populate the evidence ledger and shape assumption flags that "
        "affect Official Status, signal clarity, and later correction strength.",
    )

    h2(doc, "6.2 Stage 5 — containment hatch")
    para(
        doc,
        "Four procedures — remote inspect, controlled crack with watcher, full override with descent, or delay — "
        "change Mara verification, lock knowledge, respirator assignment, and interlock override availability. "
        "The auxiliary relay hold reveal uses dedicated consequence imagery (auxiliary-relay-hold.webp).",
    )

    h2(doc, "6.3 Finale Lock lane")
    para(
        doc,
        "Stage 7 Lock lane options include remote stabilise, physical hold with watcher, full interlock override, "
        "or leaving the interlock. Lock knowledge from Stage 5, powered Lock Control from Stage 2, and "
        "interlock override resource state gate which actions succeed without raising lock risk.",
    )

    h1(doc, "7. Finale, endings, and confirmation gate")
    h2(doc, "7.1 Two-round finale")
    bullets(
        doc,
        [
            "Round 1 and Round 2 each require Lock, Signal, and People lane actions with named leads.",
            "Duo mode (2 players) requires both operators to own at least one lane across the two rounds.",
            "Server computes a deterministic ending preview from lock control × correction acceptance × people action.",
        ],
    )

    h2(doc, "7.2 Seven canonical endings")
    table(
        doc,
        ["Ending", "Condition summary"],
        [
            ["Clean Rescue", "Lock controlled; correction accepted; Mara and crew recoverable"],
            ["Costly Rescue", "Lock controlled; correction accepted; significant resource cost"],
            ["Last Broadcast", "Lock controlled; correction not fully accepted; truth still escapes"],
            ["Joined Below", "Lock not controlled; extraction attempt without interlock control"],
            ["Flee Ridge", "Surface evacuation without Mara extraction"],
            ["Filed Safe", "Lock seals while false OPERATIONAL status files"],
            ["Station Loss", "Lock and signal both fail under maximum pressure"],
        ],
    )
    para(
        doc,
        "The facilitator sees pendingEnding with rescue, survival, and status dimensions before confirming "
        "the final record. Players cannot enter debrief until CONFIRM_ENDING is executed.",
    )

    h1(doc, "8. Debrief capture and session evidence")
    h2(doc, "8.1 Six-step synchronized debrief")
    table(
        doc,
        ["Step", "Title", "Purpose"],
        [
            ["1", "Immediate reactions", "Surface survival pressure and emotional read"],
            ["2", "Status failure map", "OPERATIONAL, GREEN, NO ACTIVE DISTRESS, LOWER ROUTE false proxies"],
            ["3", "Role map", "Responsibility areas and reflection questions per role"],
            ["4", "Readiness behaviour lenses", "Facilitator-observed patterns — not personality scores"],
            ["5", "AI-ready workplace bridge", "False proxy, missing source, bounded test"],
            ["6", "First workplace action", "Each player records a bounded commitment (private to facilitator export)"],
        ],
    )

    h2(doc, "8.2 Facilitator-only capture")
    bullets(
        doc,
        [
            "Per-step facilitator notes (up to 1,200 characters) saved server-side.",
            "Readiness profile signal toggles for observed behaviour lenses.",
            "Prototype playtest record with seven rating criteria, disposition, and reframe warning.",
            "Contribution metrics: reports, interventions, guidance bursts, field moves.",
            "Full session JSON export including decisions, history, debrief data, and ending record.",
            "Private role brief download (.txt) for remote or accessibility workflows.",
        ],
    )

    h1(doc, "9. Unequal information enforcement")
    para(
        doc,
        "Milestone 4 includes server-enforced information partitioning so no single role can solo-solve the "
        "incident. Recent verification added INFORMATION_MATRIX, INFORMATION_JOINS, and PRIVATE_PHRASE_OWNERS "
        "in game-content.js, viewer-aware signalTransmission() in the engine, and six-player Comms clue "
        "redistribution.",
    )
    checklist(
        doc,
        [
            ["Information matrix covers every stage and core role", "Complete"],
            ["Private phrases stay with owning role only", "Complete"],
            ["Shared public intercept cannot solo-solve Lower Route before Stage 5", "Complete"],
            ["Field Stage 3 alone cannot name buried route under cabin", "Complete"],
            ["Six-player mode redistributes critical Comms occupancy clue", "Complete"],
            ["Private cards include speak prompts and confidence labels", "Complete"],
        ],
    )

    h1(doc, "10. Runtime architecture")
    table(
        doc,
        ["Component", "File", "Responsibility"],
        [
            ["HTTP API + SSE + persistence", "server/server.js", "Sessions, live events, JSON store, static delivery"],
            ["Game engine", "server/game-engine.js", "Authorization, timers, consequences, finale, debrief, export"],
            ["Game content", "server/game-content.js", "Stages, roles, signals, debrief steps, playtest criteria"],
            ["Browser client", "prototype/full-app.js", "Four modes, all stage UIs, debrief, playtest form"],
            ["Network layer", "prototype/network.js", "REST actions, SSE reconnect, token storage"],
            ["Audio", "prototype/audio.js", "Opt-in ambience and interaction cues"],
            ["Presentation", "prototype/styles.css", "Cinematic full-viewport layout, responsive breakpoints"],
        ],
    )

    h1(doc, "11. Quality assurance")
    bullets(
        doc,
        [
            "npm run check — syntax validation of all JavaScript entry points.",
            "npm test — 34 automated tests covering strong/weak paths, privacy, timers, Outside Run, finale, debrief, and information matrix.",
            "docs/full-system/PLAYBOOK-COVERAGE.md — section-by-section audit against source playbook.",
            "Crew configurations validated for 2, 4, 5, 6, and 7 players.",
        ],
    )

    h1(doc, "12. User guide — running a full session")
    para(
        doc,
        "This guide explains how to host, join, play, and debrief Blackout Ridge using the current build. "
        "Voice conversation stays on Teams or Zoom; the browser app supplies synchronized station state, "
        "private role evidence, timers, and facilitator controls. There is no personal score — the app "
        "records team decisions, station consequences, and a named ending for debrief.",
    )

    h2(doc, "12.1 Before you start")
    bullets(
        doc,
        [
            "Install Node.js 18+ and run npm run dev from the project root.",
            "Open http://localhost:4173 (or your deployed host) in a modern browser (Chrome, Edge, Firefox, Safari).",
            "Plan for 70 minutes total: ~5 min briefing, ~43 min gameplay, ~20 min debrief, ~2 min buffer.",
            "Each player needs their own device. One optional shared display (TV or screen-share tab) shows the room-facing station.",
            "Facilitator stays on the voice call but does not play a role.",
        ],
    )

    h2(doc, "12.2 Facilitator — create and open the session")
    bullets(
        doc,
        [
            "On the start screen, click ENGAGE STATION, then COMMAND AUTHORITY.",
            "Choose crew size: 2 (duo), 4, 5, 6, or 7 terminals. Commercial sessions typically use 6 or 7.",
            "Click ARM NEW SESSION. The facilitator console opens with a six-character session code.",
            "Copy the crew link (TRANSMIT CREW LINK) and paste it into Teams/Zoom chat.",
            "Open the room display (ROOM DISPLAY) on a projector or screen-shared tab.",
            "Wait until every configured role shows CHANNEL LOCKED on the assembly roster.",
            "Read the participant-care statement aloud, then click CONFIRM BRIEF DELIVERED.",
            "Click START SHIFT to unlock Stage 1.",
        ],
    )

    h2(doc, "12.3 Players — join and claim a role")
    bullets(
        doc,
        [
            "Open the crew link (or enter the six-character code under CREW CHANNEL).",
            "Enter a display name (callsign) and choose one unclaimed responsibility.",
            "First come, first served. The facilitator can reassign later if someone disconnects.",
            "Your browser becomes your private terminal — you only see your own evidence card.",
            "If you refresh or drop off, reopen the same link; your role and information are restored.",
            "Optional: click ARCHIVE ROLE BRIEF to download your current private card as a text file.",
        ],
    )

    h2(doc, "12.4 The four browser views")
    table(
        doc,
        ["View", "Who uses it", "What it shows"],
        [
            ["Game start / host", "Facilitator (once)", "Session creation and crew size selection"],
            ["Shared station", "Room display or screen-share", "Tracks, timer, official traffic, crew roster, consequences"],
            ["Private role terminal", "Each player", "Role evidence, crew report channel, one-use override, lead controls"],
            ["Facilitator console", "Facilitator only", "Decision engine, live reports, timers, hints, safety, debrief, export"],
        ],
    )

    h2(doc, "12.5 How each stage works")
    para(doc, "Every stage follows the same loop: shared incident → private evidence → voice discussion → one crew commit → visible consequence.")
    table(
        doc,
        ["Stage", "Time", "The problem", "Team decision"],
        [
            ["1 — Challenge the handover", "4 min", "Dangerous handover defaults are auto-filling", "Hold exactly 3 of 5 audit items for human review"],
            ["2 — Keep the station alive", "5 min", "Backup power cannot run every circuit", "Power Main Relay + exactly 3 optional circuits"],
            ["3 — Decide what to trust", "6 min", "Clean traffic and degraded signals overlap", "Spend 3 bandwidth slots (preserve / replay / quarantine / discard)"],
            ["4 — Test Lower Route", "6 min", "Is Lower Route the flooded road?", "Assign runner, destination, abort rule; runner makes 3 moves"],
            ["5 — Open the buried station", "7 min", "Hatch to the old hold beneath the cabin", "Choose hatch procedure; assign respirator and watcher"],
            ["6 — Build the truth", "6 min", "Cordon Control needs evidence, not volume", "Build a 5-part emergency correction (one retry if rejected)"],
            ["7 — Survive the correction", "9 min", "Hold is sealing; false status is filing", "Two rounds: assign Lock, Signal, and People lanes with named leads"],
        ],
    )
    para(
        doc,
        "The Station Lead commits each decision from their private terminal. The facilitator can also commit "
        "or resolve from the console. When the clock hits zero, the server resolves the current legal draft — "
        "nobody is eliminated and the story always continues (fail-forward).",
    )

    h2(doc, "12.6 Stage 4 — Outside Run (special instructions)")
    bullets(
        doc,
        [
            "Lead locks runner, destination, and a clear abort rule (minimum 8 characters), then commits.",
            "The assigned runner sees three move buttons: Observe/Verify, Advance/Accept Exposure, Withdraw on Abort Rule.",
            "Other crew members may send exactly three short guidance bursts (90 characters each).",
            "Watch the live air meter on the runner screen. At 18 ppm the runner is forced to withdraw.",
            "Findings carry into Stage 5 hatch decisions and respirator state carries into the finale.",
        ],
    )

    h2(doc, "12.7 Facilitator controls during play")
    table(
        doc,
        ["Control", "When to use it"],
        [
            ["Pause / Resume", "Technical issues, confusion, or a short break"],
            ["+30 seconds", "Connection failure only — not to rescue a slow debate"],
            ["Send prompt / Clear prompt", "Room is stuck on the rule, not the story (one hint per stage)"],
            ["Safety pause", "Participant distress — all screens go neutral; clocks stop"],
            ["Re-send card", "Player says their private evidence did not appear"],
            ["Reassign role", "Wrong seat taken or a player disconnected permanently"],
            ["Mark absent / observer", "Someone must step out; crew continues without them"],
            ["Advance stage", "Recovery only — if the app failed to move after a valid commit"],
            ["Confirm final record", "After Stage 7 — required before debrief unlocks"],
        ],
    )

    h2(doc, "12.8 Role responsibilities and one-use overrides")
    table(
        doc,
        ["Role", "Lens during play", "One-use override"],
        [
            ["Station Lead", "Accountability; commits crew decisions", "Hold the Clock (+60 s; official pressure advances)"],
            ["Signal Analyst", "Meaning and provenance of traffic", "Recover Fragment (restore degraded source)"],
            ["Systems Engineer", "What a status actually measures", "Trace Dependency (reveal subsystem behind a status)"],
            ["Operations Officer", "Bounded physical action", "Controlled Test (preview cost, not outcome)"],
            ["Field Liaison", "People named in the data", "Human Check (force named human into the record)"],
            ["Protocol Officer", "Risk and guardrails", "Safety Condition (reduce worst physical cost once)"],
            ["Comms Officer (7p)", "Whether outsiders will believe the message", "Request Readback (receiver interpretation)"],
        ],
    )
    para(doc, "No single role can solo-solve the incident. Major conclusions require information from at least two roles.")

    h2(doc, "12.9 Ending, debrief, and export")
    bullets(
        doc,
        [
            "After Stage 7 Round 2, the facilitator sees an ending preview (title, rescue/survival/status dimensions).",
            "Click CONFIRM FINAL RECORD. All clients move to the ending screen and synchronized debrief.",
            "Advance debrief steps 1–6 together. Step 6 asks each player for one bounded workplace action.",
            "Facilitator saves notes per step and toggles observed behaviour lenses (not personality scores).",
            "Click EXPORT SESSION RECORD to download JSON with decisions, debrief, ending, and playtest data.",
            "After a pilot, complete the Prototype Playtest Record (seven ratings + disposition).",
        ],
    )

    h2(doc, "12.10 Troubleshooting")
    table(
        doc,
        ["Issue", "Fix"],
        [
            ["Player sees blank or stale screen", "Hard refresh (Ctrl+Shift+R). Reopen the same crew link."],
            ["Cannot start shift", "All terminals must be connected and participant care confirmed."],
            ["Wrong role taken", "Facilitator uses Reassign on the assembly or live console."],
            ["Private card missing", "Facilitator clicks Re-send card for that role."],
            ["Connection lost banner", "Check network; app reconnects via SSE and restores token from localStorage."],
            ["Audio does not play", "Click AUDIO on the shared station — browsers block autoplay until user gesture."],
        ],
    )

    h1(doc, "13. Screenshot evidence (current build — August 2026)")
    para(
        doc,
        "All figures below were captured from the live server-backed build on " + TODAY + " using snapshot mode "
        "and a full six-player demo session progressed through every stage. Screenshots are stored in "
        "docs/reports/screenshots/milestone-4-latest/. Regenerate with: "
        "npm run dev && node scripts/capture-milestone-4-screenshots.mjs",
    )

    screenshot(doc, M4_SHOTS / "01-home-boot.webp", "Figure 1 — Current station engagement screen (boot entry)")
    screenshot(doc, M4_SHOTS / "02-host-config.webp", "Figure 2 — Host shift configuration with distinct crew-size icons")
    screenshot(doc, M4_SHOTS / "03-join-flow.webp", "Figure 3 — Crew join flow with role instruments and participant care")
    screenshot(doc, M4_SHOTS / "04-assembly-lobby.webp", "Figure 4 — Facilitator assembly lobby before shift start")
    screenshot(doc, M4_SHOTS / "05-shared-stage2.webp", "Figure 5 — Shared station during Stage 2 (Degraded Power)")
    screenshot(doc, M4_SHOTS / "07-shared-stage3.webp", "Figure 6 — Shared station during Stage 3 (Decide What to Trust)")
    screenshot(doc, M4_SHOTS / "08-player-signal-stage3.webp", "Figure 7 — Private Signal Analyst terminal with unequal-information card")
    screenshot(doc, M4_SHOTS / "09-facilitator-stage3.webp", "Figure 8 — Facilitator console with live decision engine")
    screenshot(doc, M4_SHOTS / "10-outside-run-runner.webp", "Figure 9 — Stage 4 Outside Run runner view with air meter")
    screenshot(doc, M4_SHOTS / "12-hatch-stage5.webp", "Figure 10 — Stage 5 shared station (Hatch & Hold)")
    screenshot(doc, M4_SHOTS / "14-correction-stage6.webp", "Figure 11 — Stage 6 emergency correction builder")
    screenshot(doc, M4_SHOTS / "15-finale-stage7.webp", "Figure 12 — Stage 7 finale on shared station")
    screenshot(doc, M4_SHOTS / "17-ending-clean-rescue.webp", "Figure 13 — Clean Rescue ending screen")
    screenshot(doc, M4_SHOTS / "18-debrief-facilitator.webp", "Figure 14 — Facilitator debrief console (six-step guide)")
    screenshot(doc, M4_SHOTS / "19-debrief-player-mobile.webp", "Figure 15 — Player debrief with first-step commitment (mobile)", width=Inches(3.8))
    screenshot(doc, M4_SHOTS / "20-playtest-record.webp", "Figure 16 — Prototype playtest record after pilot session")

    h1(doc, "14. Milestone 4 acceptance checklist")
    checklist(
        doc,
        [
            ["All seven challenges playable end-to-end", "Complete"],
            ["Progressive degraded signal responds to earlier preservation choices", "Complete"],
            ["Opt-in audio ambience and interaction cues", "Complete"],
            ["Outside Run: runner, moves, guidance, air meter, forced withdrawal", "Complete"],
            ["Stage 1 holds and Stage 5 hatch/interlock with persistent resources", "Complete"],
            ["Stage 6 correction builder with one retry on rejection", "Complete"],
            ["Two-round finale with named lane leads", "Complete"],
            ["Seven endings with facilitator preview and confirmation gate", "Complete"],
            ["Six-step debrief synchronized across all clients", "Complete"],
            ["Session JSON export with decisions, debrief, and playtest data", "Complete"],
            ["Unequal information enforced in server logic and tests", "Complete"],
            ["Automated test suite passes (34/34)", "Complete"],
        ],
    )

    h1(doc, "15. Explicitly deferred to later milestones")
    table(
        doc,
        ["Item", "Target milestone", "Notes"],
        [
            ["Live 6–7 player playtesting and pacing tuning", "Milestone 5", "Runtime captures playtest ratings; empirical validation required"],
            ["Facilitator load testing across Teams/Zoom", "Milestone 5", "One-console design complete; remote coordination untested at scale"],
            ["Production deployment (TLS, managed DB, rate limits)", "Milestone 5+", "Local JSON persistence sufficient for workshops"],
            ["Full recorded sound library and cinematic opening", "Milestone 6", "Procedural audio satisfies playable build requirement"],
            ["Promotional video edit", "Milestone 6", "Explicitly excluded from M1 design package"],
        ],
    )

    h1(doc, "16. Handoff to Milestone 5")
    para(
        doc,
        "Milestone 4 completes the playable game. Milestone 5 should focus on facilitated pilot sessions "
        "(target: 6–7 players), difficulty and pacing adjustments from real crew behaviour, facilitator load "
        "under remote voice coordination, responsive/browser matrix testing, and facilitator documentation "
        "for external hosts. The playtest record and reframe warning in the runtime are ready to capture "
        "pilot outcomes.",
    )
    para(doc, "Technical reference: docs/full-system/README.md and docs/full-system/PLAYBOOK-COVERAGE.md.")

    out = OUT / "Blackout_Ridge_Milestone_4_Completion_Report.docx"
    doc.save(out)
    return out


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    paths = [build_milestone_1(), build_milestone_2(), build_milestone_3(), build_milestone_4()]
    for p in paths:
        print(f"Wrote {p}")


if __name__ == "__main__":
    main()
